import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Laporan_Pendahuluan_SIMPAH.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Laporan_Pendahuluan_SIMPAH_v1.docx"
LOGO_PATH = r"u:\Project\simpah-rilis v1\logo.png"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # #059669 Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # #1F2937 Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # #6B7280 Muted Gray
COLOR_LIGHT_BG = "F0FDF4"                 # #F0FDF4 Emerald Tint
COLOR_BORDER = "10B981"                   # #10B981 Emerald
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_p(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def create_document():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # COVER PAGE
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(30)
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
        
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_after = Pt(12)
    r_badge = p_badge.add_run("CV. EXADATA — DOKUMEN RESMI PENGEMBANGAN SOFTWARE")
    r_badge.font.name = 'Arial'
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_PRIMARY
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("LAPORAN PENDAHULUAN\nPENGEMBANGAN SOFTWARE APLIKASI PELAYANAN PERSAMPAHAN")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_DARK
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(35)
    r_sub = p_sub.add_run("Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) Terintegrasi\nKabupaten Banjarnegara, Jawa Tengah")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    
    div_table = doc.add_table(rows=1, cols=1)
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_background(div_cell, HEX_PRIMARY)
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    p_div = div_cell.paragraphs[0]
    p_div.paragraph_format.space_after = Pt(0)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(35)
    p_meta.paragraph_format.space_after = Pt(4)
    r_m1 = p_meta.add_run("Instansi Pemilik Pekerjaan: ")
    r_m1.font.bold = True
    p_meta.add_run("Dinas Perumahan, Kawasan Permukiman dan Lingkungan Hidup (DPPKPLH) Kabupaten Banjarnegara\n")
    
    p_meta2 = doc.add_paragraph()
    p_meta2.paragraph_format.space_after = Pt(4)
    r_m2 = p_meta2.add_run("Penyedia Jasa / Pelaksana: ")
    r_m2.font.bold = True
    p_meta2.add_run("CV. EXADATA\n")
    
    p_meta3 = doc.add_paragraph()
    p_meta3.paragraph_format.space_after = Pt(4)
    r_m3 = p_meta3.add_run("Durasi Pelaksanaan: ")
    r_m3.font.bold = True
    p_meta3.add_run("30 Hari Kerja (1 – 30 Juli 2026)\n")
    
    p_meta4 = doc.add_paragraph()
    p_meta4.paragraph_format.space_after = Pt(0)
    r_m4 = p_meta4.add_run("Tahun Anggaran: ")
    r_m4.font.bold = True
    p_meta4.add_run("2026")

    doc.add_page_break()
    
    body_section = doc.sections[-1]
    header = body_section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("CV. EXADATA — Laporan Pendahuluan Pengembangan Software SIMPAH")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = COLOR_MUTED
    
    footer = body_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("DPPKPLH Kabupaten Banjarnegara | Laporan Pendahuluan Pekerjaan")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = COLOR_MUTED

    # BAB 1
    add_heading_1(doc, "BAB 1 — PENDAHULUAN")
    add_heading_2(doc, "1.1. Latar Belakang")
    add_body_p(doc, "Pengelolaan sampah urban di era modern membutuhkan pendekatan berbasis data yang presisi, cepat, dan transparan. Kabupaten Banjarnegara, sebagai salah satu wilayah berkembang di Provinsi Jawa Tengah, memiliki dinamika pertumbuhan timbulan sampah yang terus meningkat seiring bertambahnya populasi dan aktivitas ekonomi masyarakat.")
    add_body_p(doc, "Selama ini, proses pencatatan dan monitoring volume sampah dari berbagai sumber (rumah tangga, pasar, TPS, hingga TPA Winong) masih dilakukan secara manual dan terpisah-pisah. Hal ini menimbulkan berbagai potensi masalah strategis, antara lain risiko manipulasi data tonase, keterlambatan pelaporan berkala, ketidakakuratan data pemilahan, serta kesulitan pimpinan daerah dalam mengevaluasi efektivitas alokasi armada dan anggaran pengelolaan lingkungan hidup.")
    add_body_p(doc, "Sejalan dengan komitmen Dinas Perumahan, Kawasan Permukiman dan Lingkungan Hidup (DPPKPLH) Kabupaten Banjarnegara dalam mewujudkan tata kelola kebersihan berbasis 'Smart Waste Intelligence', diperlukan sebuah perangkat lunak terintegrasi yang mampu memonitor seluruh rantai pasok pengelolaan sampah secara real-time. Untuk menjawab tantangan tersebut, CV. EXADATA dipercaya untuk melaksanakan kegiatan Pekerjaan Pengembangan Software Aplikasi Pelayanan Persampahan — Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH).")

    add_heading_2(doc, "1.2. Maksud dan Tujuan Pekerjaan")
    add_body_p(doc, "Maksud dari pelaksanaan pekerjaan ini adalah membangun dan mengimplementasikan aplikasi software SIMPAH terintegrasi berbasis Progressive Web App (PWA) yang menghubungkan seluruh pemangku kepentingan pengelolaan sampah di Kabupaten Banjarnegara.")
    add_body_p(doc, "Tujuan yang ingin dicapai dari pengembangan software ini meliputi:")
    add_bullet_p(doc, "Mendigitalisasi seluruh rantai pencatatan volume dan kategori sampah dari tingkat warga/RT/RW, fasilitas TPS3R/Bank Sampah, armada pengangkut, hingga TPA Winong.", "1. Akurasi & Digitalisasi Data: ")
    add_bullet_p(doc, "Menyediakan aplikasi berbasis PWA dengan arsitektur Offline-First, yang memungkinkan petugas di lokasi minim sinyal tetap dapat menginput data secara stabil di lapangan.", "2. Kemudahan Aksesibilitas Lapangan: ")
    add_bullet_p(doc, "Menerapkan mekanisme validasi data berjenjang dari Koordinator Lapangan sebelum data masuk ke kalkulasi resmi dashboard.", "3. Verifikasi & Integritas Data: ")
    add_bullet_p(doc, "Menyajikan Dashboard Ringkasan Eksekutif untuk Kepala Dinas dan Pimpinan Daerah yang dilengkapi pemodelan prediktif Machine Learning (Meta Prophet) untuk memproyeksikan timbulan sampah 7 hari ke depan.", "4. Decision Support System (DSS): ")
    add_bullet_p(doc, "Menyediakan modul laporan otomatis yang terpetakan sesuai format standar Pelaporan Informasi Pengelolaan Sampah Nasional (SIPSN) Kementerian LHK.", "5. Pemenuhan Standar SIPSN KLHK: ")
    add_bullet_p(doc, "Memudahkan masyarakat dalam menyampaikan laporan pengaduan kebersihan berbasis koordinat GPS dan foto bukti.", "6. Pelayanan Publik & Kanal Pengaduan: ")

    add_heading_2(doc, "1.3. Sasaran Strategis Pekerjaan")
    add_body_p(doc, "Sasaran strategis yang hendak diwujudkan melalui pengembangan software SIMPAH adalah:")
    add_bullet_p(doc, "Terbangunnya platform PWA SIMPAH yang responsif, aman, dan siap digunakan oleh 4 tingkatan peran pengguna (Warga, Petugas Lapangan, Eksekutif, dan Admin).")
    add_bullet_p(doc, "Tersedianya database PostgreSQL terpusat di Supabase dengan enkripsi aman dan keamanan data tingkat baris (Row Level Security / RLS).")
    add_bullet_p(doc, "Tersambungnya 278 Desa/Kelurahan di 28 Kecamatan Kabupaten Banjarnegara ke dalam satu ekosistem monitoring persampahan daerah.")
    add_bullet_p(doc, "Terlaksananya pelatihan (training) dan pendampingan teknis bagi petugas lapangan dan administrator dinas.")

    add_heading_2(doc, "1.4. Landasan Hukum & Kebijakan")
    add_body_p(doc, "Pengembangan software SIMPAH mengacu pada regulasi dan kebijakan perundang-undangan berikut:")
    add_bullet_p(doc, "Undang-Undang Nomor 18 Tahun 2008 tentang Pengelolaan Sampah.", "1. ")
    add_bullet_p(doc, "Peraturan Pemerintah Nomor 81 Tahun 2012 tentang Pengelolaan Sampah Rumah Tangga dan Sampah Sejenis Sampah Rumah Tangga.", "2. ")
    add_bullet_p(doc, "Peraturan Presiden Nomor 97 Tahun 2017 tentang Kebijakan dan Strategi Nasional Pengelolaan Sampah Rumah Tangga dan Sampah Sejenis Sampah Rumah Tangga (Jakstranas).", "3. ")
    add_bullet_p(doc, "Peraturan Daerah Kabupaten Banjarnegara tentang Pengelolaan Sampah dan Kebersihan Lingkungan Hidup.", "4. ")
    add_bullet_p(doc, "Pedoman Sistem Informasi Pengelolaan Sampah Nasional (SIPSN) Kementerian Lingkungan Hidup dan Kehutanan (KLHK).", "5. ")

    add_heading_2(doc, "1.5. Lokasi & Wilayah Implementasi Pekerjaan")
    add_body_p(doc, "Lokasi pelaksanaan pekerjaan pengembangan software dilakukan di kantor CV. EXADATA dan DPPKPLH Kabupaten Banjarnegara. Wilayah cakupan implementasi sistem meliputi seluruh unit operasional kebersihan di Kabupaten Banjarnegara, Jawa Tengah (28 Kecamatan, 278 Desa/Kelurahan, TPS3R, Bank Sampah, dan TPA Winong).")

    add_heading_2(doc, "1.6. Sistematika Penulisan Laporan Pendahuluan")
    add_body_p(doc, "Laporan Pendahuluan ini disusun dengan sistematika penulisan sebagai berikut:")
    add_bullet_p(doc, "Memuat latar belakang, maksud dan tujuan, sasaran strategis, landasan hukum, lokasi pekerjaan, dan sistematika laporan.", "BAB 1 — PENDAHULUAN: ")
    add_bullet_p(doc, "Memuat profil pengelolaan sampah daerah, analisis permasalahan eksisting, urgensi software, dan kaitan dengan standar SIPSN.", "BAB 2 — GAMBARAN UMUM & KONDISI EKSISTING: ")
    add_bullet_p(doc, "Memuat pendekatan Agile/Scrum, arsitektur teknis PWA, Supabase PostgreSQL, ML Prophet Forecasting, AI Assistant, dan RBAC.", "BAB 3 — METODOLOGI & ARSITEKTUR TEKNIS SOFTWARE: ")
    add_bullet_p(doc, "Memuat rincian fitur/modul aplikasi, tahapan kerja, Gantt Chart jadwal kerja 30 Hari Kerja (1-30 Juli 2026), serta mitigasi risiko.", "BAB 4 — RUANG LINGKUP & RENCANA KERJA: ")
    add_bullet_p(doc, "Memuat struktur organisasi tim CV. EXADATA, daftar kualifikasi tenaga ahli, dan alokasi person-month.", "BAB 5 — STRUKTUR ORGANISASI & TENAGA AHLI: ")
    add_bullet_p(doc, "Memuat daftar produk keluaran software, rencana pelatihan, serta kesimpulan tahap pendahuluan.", "BAB 6 — RENCANA KELUARAN & PENUTUP: ")

    # BAB 2
    add_heading_1(doc, "BAB 2 — GAMBARAN UMUM & KONDISI EKSISTING")
    add_heading_2(doc, "2.1. Profil Pengelolaan Sampah Kabupaten Banjarnegara")
    add_body_p(doc, "Kabupaten Banjarnegara memiliki wilayah geografis yang membentang dari dataran rendah hingga kawasan pegunungan (Dieng). Pengelolaan sampah pelayanan dinas mencakup wilayah pemukiman padat, fasilitas umum/pasar, hingga wilayah perdesaan.")
    add_body_p(doc, "Rantai pasok pengelolaan sampah eksisting terdiri dari sumber sampah, TPS, TPS3R & Bank Sampah, armada pengangkut, dan TPA Winong.")

    add_heading_2(doc, "2.2. Analisis Permasalahan & Kendala Operasional Eksisting")
    add_body_p(doc, "Berdasarkan hasil studi pendahuluan, ditemukan beberapa permasalahan utama dalam sistem operasional eksisting:")

    add_callout_box(
        doc,
        "IDENTIFIKASI PERMASALAHAN EKSISTING (BEFORE SIMPAH)",
        [
            "Pencatatan Manual Berbasis Kertas: Data tonase sampah di TPS dan ritase pengangkutan masih dicatat pada formulir kertas, rawan rusak, hilang, dan mengalami kesalahan input (human error).",
            "Risiko Manipulasi Data: Tidak adanya bukti foto timbangan dan koordinat GPS membuat validasi tonase sampah sulit diverifikasi secara obyektif.",
            "Kendala Sinyal Internet di Wilayah Pelosok: Petugas di wilayah pegunungan sering mengalami kegagalan input data saat menggunakan web aplikasi biasa yang membutuhkan koneksi stabil.",
            "Keterlambatan Konsolidasi Laporan: Pengumpulan data dari 28 kecamatan membutuhkan waktu berminggu-minggu, sehingga Laporan SIPSN KLHK sering terlambat disusun.",
            "Penanganan Pengaduan Kurang Terukur: Aduan warga mengenai penumpukan sampah liar sering tidak terdata dengan nomor resi resmi, sehingga status penanganannya sulit dilacak."
        ],
        bg_hex="FEF2F2",
        border_hex="EF4444"
    )

    add_heading_2(doc, "2.3. Urgensi Pengadaan Software SIMPAH (Offline-First PWA)")
    add_body_p(doc, "Transformasi digital melalui pengadaan software SIMPAH berbasis Progressive Web App (PWA) dengan Offline-First Engine (IndexedDB & auto-sync) dan Validasi Foto GPS.")

    add_heading_2(doc, "2.4. Kebijakan Pelaporan & Integrasi Standar SIPSN KLHK")
    add_body_p(doc, "Pemetaan otomatis 8 kategori sampah standar SIPSN (Organik, Plastik, Kertas, Logam, Kaca, Karet, Kain, Lainnya) untuk kemudahan ekspor laporan resmi.")

    # -------------------------------------------------------------
    # BAB 3 — METODOLOGI & ARSITEKTUR TEKNIS SOFTWARE
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 3 — METODOLOGI & ARSITEKTUR TEKNIS SOFTWARE")
    
    add_heading_2(doc, "3.1. Metodologi Pengembangan Software (Agile / Scrum)")
    add_body_p(doc, "CV. EXADATA menerapkan metodologi Agile Development Lifecycle dengan kerangka kerja Scrum untuk memastikan fleksibilitas, kecepatan rilis, dan kolaborasi intensif dengan pihak DPPKPLH Kabupaten Banjarnegara. Siklus pengembangan dibagi ke dalam 4 Sprint (mingguan) selama 30 Hari Kerja (1 – 30 Juli 2026).")

    add_heading_2(doc, "3.2. Arsitektur Sistem Terintegrasi")
    add_body_p(doc, "Arsitektur software SIMPAH mengombinasikan teknologi web modern yang andal dan scalable:")
    add_bullet_p(doc, "HTML5, Vanilla CSS (Glassmorphism design system), JavaScript ES6 modern, dan Vite bundler.", "Frontend PWA Engine: ")
    add_bullet_p(doc, "PostgreSQL terenkripsi dengan otentikasi Supabase Auth & Row Level Security (RLS).", "Database & Auth (Supabase): ")
    add_bullet_p(doc, "IndexedDB + Service Worker untuk penyimpanan lokal di HP dan sinkronisasi otomatis via `initSync()`.", "Offline-First Engine: ")
    add_bullet_p(doc, "Backend Python (FastAPI + Meta Prophet) untuk proyeksi volume sampah 7 hari ke depan.", "Machine Learning Backend: ")
    add_bullet_p(doc, "Integration AI Chatbot menggunakan LLM Qwen (Alibaba Cloud MaaS) untuk asisten virtual SIMPAH Buddy.", "AI Assistant (LLM Qwen): ")

    add_heading_2(doc, "3.3. Matriks Keamanan & Hak Akses Berbasis Peran (RBAC)")
    add_body_p(doc, "Keamanan data dikendalikan melalui Role-Based Access Control (RBAC) yang membagi hak akses ke dalam 4 peran utama (Warga, Petugas, Eksekutif, Admin) dan 4 sub-peran petugas (Kader, Operator TPS3R, Angkut, Koordinator).")

    # -------------------------------------------------------------
    # BAB 4 — RUANG LINGKUP & RENCANA KERJA
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 4 — RUANG LINGKUP & RENCANA KERJA")
    
    add_heading_2(doc, "4.1. Ruang Lingkup Modul & Fitur Software")
    add_body_p(doc, "Ruang lingkup pengembangan mencakup 6 modul utama: Modul Portal Publik, Modul PWA Petugas Lapangan, Modul Validasi Koordinator, Modul Dashboard Eksekutif & ML Forecasting, Modul Super Admin & Master Data, serta Modul AI Assistant SIMPAH Buddy.")

    add_heading_2(doc, "4.2. Jadwal Pelaksanaan Pekerjaan (Gantt Chart 30 Hari Kerja)")
    add_body_p(doc, "Pekerjaan dilaksanakan selama 30 Hari Kerja terhitung dari tanggal 1 Juli 2026 sampai dengan 30 Juli 2026 dengan rincian tahapan:")

    headers_gantt = ["Tahapan Kegiatan", "Minggu 1", "Minggu 2", "Minggu 3", "Minggu 4"]
    data_gantt = [
        ["Tahap Inisiasi & Laporan Pendahuluan", "✓", "—", "—", "—"],
        ["Analisis Kebutuhan & Desain Arsitektur/DB", "✓", "✓", "—", "—"],
        ["Pengembangan PWA Lapangan & Offline Engine", "—", "✓", "✓", "—"],
        ["Pengembangan Dashboard Eksekutif & ML Prophet", "—", "—", "✓", "✓"],
        ["Pengujian Internal & Laporan Antara", "—", "—", "✓", "—"],
        ["UAT, Serah Terima & Laporan Akhir", "—", "—", "—", "✓"]
    ]
    
    table_g = doc.add_table(rows=len(data_gantt) + 1, cols=5)
    table_g.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_g = table_g.rows[0].cells
    widths_g = [Inches(3.5), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7)]
    for i, title in enumerate(headers_gantt):
        hdr_cells_g[i].text = title
        hdr_cells_g[i].width = widths_g[i]
        set_cell_background(hdr_cells_g[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells_g[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_g[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data_gantt):
        row_cells = table_g.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_g[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if cell_value == "✓":
                    run.font.bold = True
                    run.font.color.rgb = COLOR_PRIMARY

    # -------------------------------------------------------------
    # BAB 5 — STRUKTUR ORGANISASI & TENAGA AHLI
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 5 — STRUKTUR ORGANISASI & TENAGA AHLI")
    
    add_heading_2(doc, "5.1. Struktur Tim Pelaksana Pekerjaan (CV. EXADATA)")
    add_body_p(doc, "Untuk menjamin kualitas dan ketepatan waktu penyelesaian software SIMPAH, CV. EXADATA menurunkan tim tenaga ahli berpengalaman di bidang pengembangan perangkat lunak pemerintah:")

    headers_team = ["No", "Nama / Peran Tenaga Ahli", "Kualifikasi Keahlian", "Alokasi"]
    data_team = [
        ["1", "Project Manager / Team Leader", "S2/S1 Teknik Informatika / Manajemen Proyek Software (Pengalaman 8+ th)", "1.0 PM"],
        ["2", "Senior System Analyst & Architect", "S1 Teknik Informatika / Software Architecture (Pengalaman 6+ th)", "1.0 PM"],
        ["3", "Lead Frontend PWA Engineer", "S1 Ilmu Komputer / PWA & Web Specialist (Pengalaman 5+ th)", "1.0 PM"],
        ["4", "Database & Backend Specialist", "S1 Teknik Informatika / PostgreSQL & Supabase Specialist (Pengalaman 5+ th)", "1.0 PM"],
        ["5", "Data Scientist / ML Engineer", "S1 Matematika/Data Science / Python Prophet Specialist (Pengalaman 4+ th)", "0.5 PM"],
        ["6", "UI/UX Designer", "S1 Desain / UI/UX Glassmorphic Specialist (Pengalaman 4+ th)", "0.5 PM"],
        ["7", "QA & Security Tester", "S1 Teknik Informatika / Software Testing & Security (Pengalaman 4+ th)", "0.5 PM"]
    ]

    table_t = doc.add_table(rows=len(data_team) + 1, cols=4)
    table_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_t = table_t.rows[0].cells
    widths_t = [Inches(0.5), Inches(2.2), Inches(3.0), Inches(0.8)]
    for i, title in enumerate(headers_team):
        hdr_cells_t[i].text = title
        hdr_cells_t[i].width = widths_t[i]
        set_cell_background(hdr_cells_t[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells_t[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_t[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data_team):
        row_cells = table_t.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_t[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)

    # -------------------------------------------------------------
    # BAB 6 — RENCANA KELUARAN & PENUTUP
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 6 — RENCANA KELUARAN (DELIVERABLES) & PENUTUP")
    
    add_heading_2(doc, "6.1. Rencana Produk Output Pekerjaan")
    add_body_p(doc, "Keluaran pekerjaan yang akan diserahterimakan kepada DPPKPLH Kabupaten Banjarnegara meliputi:")
    add_bullet_p(doc, "Aplikasi web PWA SIMPAH terinstall di server production siap pakai.", "1. Software PWA SIMPAH Production: ")
    add_bullet_p(doc, "Kode sumber aplikasi lengkap beserta repositori deployment.", "2. Source Code & Repository: ")
    add_bullet_p(doc, "Struktur database PostgreSQL di Supabase beserta skema RLS.", "3. Database Schema & Migration SQL: ")
    add_bullet_p(doc, "Buku Panduan Operasional Aplikasi SIMPAH dalam format Word (.docx) & PDF.", "4. Manual Book / Panduan Operasional: ")
    add_bullet_p(doc, "Dokumen resmi Laporan Pendahuluan, Laporan Antara, dan Laporan Akhir.", "5. Dokumen Laporan Pekerjaan: ")

    add_heading_2(doc, "6.2. Kesimpulan Tahap Pendahuluan & Langkah Selanjutnya")
    add_body_p(doc, "Tahap pendahuluan ini menegaskan kesiapan penuh CV. EXADATA dalam melaksanakan pengembangan software SIMPAH sesuai target waktu 30 Hari Kerja (1-30 Juli 2026). Langkah selanjutnya adalah pemantapan analisis kebutuhan detail dan penyusunan prototipe software untuk disajikan pada Laporan Antara.")

    # Save document safely
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully created at primary path: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary file locked by Word. Document successfully saved at fallback path: {target_out}")

if __name__ == "__main__":
    create_document()
