import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_surat_lisensi():
    doc = docx.Document()
    
    # Page setup - Margins 2.54 cm / 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    COLOR_PRIMARY = RGBColor(5, 150, 105)   # #059669 Emerald Green
    COLOR_DARK = RGBColor(31, 41, 55)      # #1F2937 Dark Slate
    COLOR_MUTED = RGBColor(107, 114, 128)  # #6B7280 Gray
    
    # Helper styles
    def style_run(run, font_name="Arial", size_pt=11, bold=False, italic=False, color=COLOR_DARK):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r1 = p_title.add_run("SURAT PERNYATAAN PEMBERIAN LISENSI")
    style_run(r1, size_pt=14, bold=True, color=COLOR_PRIMARY)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)
    r2 = p_sub.add_run("PENGGUNAAN APLIKASI SIMPAH (SISTEM INFORMASI MONITORING PENGELOLAAN SAMPAH)")
    style_run(r2, size_pt=11, bold=True, color=COLOR_DARK)
    
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_after = Pt(18)
    r_num = p_num.add_run("Nomor: 042/LIC-SIMPAH/VIII/2026")
    style_run(r_num, size_pt=10, italic=True, color=COLOR_MUTED)

    # Horizontal Divider
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(14)
    p_hr_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="059669"/></w:pBdr>')
    p_hr._p.get_or_add_pPr().append(p_hr_border)

    # Opening text
    p_open = doc.add_paragraph()
    p_open.paragraph_format.space_after = Pt(10)
    p_open.paragraph_format.line_spacing = 1.15
    r_open = p_open.add_run("Yang bertanda tangan di bawah ini:")
    style_run(r_open, size_pt=11)

    # Table for Party 1
    t1 = doc.add_table(rows=4, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    
    party1_data = [
        ("Nama Pihak Pertama", ":", "[Nama Pengembang / Perwakilan Pengembang]"),
        ("Jabatan", ":", "[Pimpinan / Lead Developer / Tim Pengembang]"),
        ("Instansi / Perusahaan", ":", "[Nama Perusahaan / Tim Pengembang SIMPAH]"),
        ("Alamat / Kontak", ":", "[Alamat Lengkap & Email / Nomor Telepon]"),
    ]
    
    col_widths = [Inches(1.8), Inches(0.2), Inches(4.5)]
    for row_idx, data in enumerate(party1_data):
        row = t1.rows[row_idx]
        for col_idx in range(3):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            run = cell.paragraphs[0].add_run(data[col_idx])
            style_run(run, size_pt=10.5, bold=(col_idx == 0))

    # Text Pihak Pertama
    p_p1 = doc.add_paragraph()
    p_p1.paragraph_format.space_before = Pt(8)
    p_p1.paragraph_format.space_after = Pt(10)
    r_p1 = p_p1.add_run("Dalam hal ini bertindak untuk dan atas nama pengembang/penyedia aplikasi SIMPAH, selanjutnya disebut sebagai PIHAK PERTAMA (Pemberi Lisensi).")
    style_run(r_p1, size_pt=11, italic=True)

    # Text Memberikan Lisensi Kepada
    p_to = doc.add_paragraph()
    p_to.paragraph_format.space_after = Pt(10)
    r_to = p_to.add_run("Dengan ini menyatakan memberikan lisensi penggunaan aplikasi kepada:")
    style_run(r_to, size_pt=11)

    # Table for Party 2
    t2 = doc.add_table(rows=4, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    
    party2_data = [
        ("Nama Instansi / Lembaga", ":", "Dinas Perumahan, Kawasan Permukiman dan Lingkungan Hidup (DPPKPLH) Kabupaten Banjarnegara"),
        ("Perwakilan / Kepala Dinas", ":", "[Nama Kepala Dinas / Pejabat Berwenang]"),
        ("NIP", ":", "[NIP Pejabat]"),
        ("Alamat Instansi", ":", "Jl. Selamanik No. 8, Banjarnegara, Jawa Tengah"),
    ]
    
    for row_idx, data in enumerate(party2_data):
        row = t2.rows[row_idx]
        for col_idx in range(3):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            run = cell.paragraphs[0].add_run(data[col_idx])
            style_run(run, size_pt=10.5, bold=(col_idx == 0))

    # Text Pihak Kedua
    p_p2 = doc.add_paragraph()
    p_p2.paragraph_format.space_before = Pt(8)
    p_p2.paragraph_format.space_after = Pt(14)
    r_p2 = p_p2.add_run("Dalam hal ini bertindak untuk dan atas nama pengguna/instansi penerima, selanjutnya disebut sebagai PIHAK KEDUA (Penerima Lisensi).")
    style_run(r_p2, size_pt=11, italic=True)

    # Statement Intro
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.paragraph_format.line_spacing = 1.15
    r_intro = p_intro.add_run("PIHAK PERTAMA memberikan Hak Lisensi Penggunaan Perangkat Lunak kepada PIHAK KEDUA dengan syarat dan ketentuan sebagai berikut:")
    style_run(r_intro, size_pt=11)

    # Terms Clauses
    clauses = [
        ("PASAL 1: OBJEK LISENSI", 
         "1. Objek lisensi adalah Hak Penggunaan Aplikasi Sistem Informasi Monitoring Pengelolaan Sampah (SIMPAH) Versi 1.2.\n"
         "2. Lisensi mencakup seluruh modul aplikasi yaitu Modul PWA Mobile Lapangan (Offline-first), Dashboard GIS Geospasial, Dashboard Eksekutif DLHK, Analitik Kewilayahan, dan Portal Aduan Publik."),
        
        ("PASAL 2: JANGKA WAKTU LISENSI", 
         "1. Lisensi penggunaan aplikasi ini diberikan untuk jangka waktu 1 (satu) Tahun Terhitung.\n"
         "2. Masa berlaku lisensi dimulai sejak tanggal 01 Agustus 2026 sampai dengan tanggal 31 Juli 2027.\n"
         "3. Lisensi dapat diperpanjang kembali setelah masa berlaku berakhir berdasarkan kesepakatan tertulis kedua belah pihak."),
        
        ("PASAL 3: SIFAT LISENSI", 
         "1. Lisensi ini bersifat Non-Eksklusif (Non-Exclusive) dan Tidak Dapat Dipindahtangankan (Non-Transferable).\n"
         "2. PIHAK KEDUA berhak menggunakan aplikasi SIMPAH secara penuh untuk kepentingan operasional pengelolaan sampah di wilayah Kabupaten Banjarnegara."),
        
        ("PASAL 4: HAK & DUKUNGAN TEKNIS PIHAK KEDUA", 
         "1. PIHAK KEDUA berhak mendapatkan akses ke seluruh fitur aplikasi sesuai dengan hak akses (Role-Based Access Control) yang telah ditentukan.\n"
         "2. PIHAK PERTAMA menyediakan pemeliharaan rutin (maintenance), perbaikan bug (bug fixes), serta pembaruan sistem keamanan selama 1 (satu) tahun masa lisensi.\n"
         "3. PIHAK PERTAMA menyediakan dukungan teknis (technical support) pada jam kerja operasional untuk menangani kendala teknis aplikasi."),
        
        ("PASAL 5: HAK KEKAYAAN INTELEKTUAL & BATASAN PENGGUNAAN", 
         "1. Hak Cipta dan Hak Kekayaan Intelektual (HKI) atas kode sumber (source code), desain, dan arsitektur aplikasi SIMPAH tetap berada sepenuhnya pada PIHAK PERTAMA.\n"
         "2. PIHAK KEDUA dilarang menggandakan, memperjualbelikan, menyewakan, memodifikasi kode sumber utama, atau memberikan lisensi turunan (sub-license) aplikasi SIMPAH kepada pihak ketiga tanpa izin tertulis dari PIHAK PERTAMA."),
        
        ("PASAL 6: KERAHASIAAN DAN KEAMANAN DATA", 
         "1. Seluruh data operasional, rekam timbangan sampah, dan data pengguna yang tersimpan di dalam database SIMPAH adalah hak milik penuh PIHAK KEDUA.\n"
         "2. PIHAK PERTAMA berkewajiban menjaga kerahasiaan dan keamanan data PIHAK KEDUA serta tidak menyalahgunakan data tersebut untuk kepentingan di luar operasional sistem.")
    ]

    for title, content in clauses:
        p_chead = doc.add_paragraph()
        p_chead.paragraph_format.space_before = Pt(10)
        p_chead.paragraph_format.space_after = Pt(3)
        p_chead.paragraph_format.keep_with_next = True
        r_ch = p_chead.add_run(title)
        style_run(r_ch, size_pt=11, bold=True, color=COLOR_PRIMARY)
        
        p_cbody = doc.add_paragraph()
        p_cbody.paragraph_format.space_after = Pt(8)
        p_cbody.paragraph_format.line_spacing = 1.15
        r_cb = p_cbody.add_run(content)
        style_run(r_cb, size_pt=10.5, color=COLOR_DARK)

    # Closing Statement
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(12)
    p_close.paragraph_format.space_after = Pt(20)
    p_close.paragraph_format.line_spacing = 1.15
    r_close = p_close.add_run(
        "Demikian Surat Pernyataan Pemberian Lisensi ini dibuat dengan sebenarnya dalam 2 (dua) rangkap bermaterai cukup "
        "dan mempunyai kekuatan hukum yang sama untuk dipergunakan sebagaimana mestinya."
    )
    style_run(r_close, size_pt=10.5)

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_after = Pt(14)
    r_date = p_date.add_run("Banjarnegara, 29 Juli 2026")
    style_run(r_date, size_pt=10.5)

    # Signature Table
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    cell_left = sig_table.rows[0].cells[0]
    cell_right = sig_table.rows[0].cells[1]
    
    cell_left.width = Inches(3.2)
    cell_right.width = Inches(3.2)
    
    # Left Signature (PIHAK KEDUA / Mengetahui)
    p_l1 = cell_left.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l1.paragraph_format.space_after = Pt(2)
    r_l1 = p_l1.add_run("PIHAK KEDUA\n(Penerima Lisensi)")
    style_run(r_l1, size_pt=10.5, bold=True)
    
    p_l2 = cell_left.add_paragraph()
    p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l2.paragraph_format.space_after = Pt(50)  # Space for signature
    r_l2 = p_l2.add_run("DPPKPLH Kabupaten Banjarnegara")
    style_run(r_l2, size_pt=9.5, italic=True, color=COLOR_MUTED)
    
    p_l3 = cell_left.add_paragraph()
    p_l3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l3.paragraph_format.space_after = Pt(2)
    r_l3 = p_l3.add_run("( ___________________________ )")
    style_run(r_l3, size_pt=10.5, bold=True)
    
    p_l4 = cell_left.add_paragraph()
    p_l4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l4 = p_l4.add_run("NIP. ....................................")
    style_run(r_l4, size_pt=9.5, color=COLOR_MUTED)

    # Right Signature (PIHAK PERTAMA / Pemberi Lisensi)
    p_r1 = cell_right.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r1.paragraph_format.space_after = Pt(2)
    r_r1 = p_r1.add_run("PIHAK PERTAMA\n(Pemberi Lisensi)")
    style_run(r_r1, size_pt=10.5, bold=True)
    
    p_r2 = cell_right.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.paragraph_format.space_after = Pt(50)  # Space for signature & stamp
    r_r2 = p_r2.add_run("Tim Pengembang SIMPAH\n[Materai Rp10.000]")
    style_run(r_r2, size_pt=9.5, italic=True, color=COLOR_MUTED)
    
    p_r3 = cell_right.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r3.paragraph_format.space_after = Pt(2)
    r_r3 = p_r3.add_run("( ___________________________ )")
    style_run(r_r3, size_pt=10.5, bold=True)
    
    p_r4 = cell_right.add_paragraph()
    p_r4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r4 = p_r4.add_run("[Jabatan / Pimpinan]")
    style_run(r_r4, size_pt=9.5, color=COLOR_MUTED)

    output_path = r"u:\Project\simpah-rilis v1\Surat_Pernyataan_Lisensi_SIMPAH.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_surat_lisensi()
