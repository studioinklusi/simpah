import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Laporan_Antara_SIMPAH.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Laporan_Antara_SIMPAH_v1.docx"
LOGO_PATH = r"u:\Project\simpah-rilis v1\logo.png"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # #059669 Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # #1F2937 Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # #6B7280 Muted Gray
COLOR_LIGHT_BG = "F0FDF4"                 # #F0FDF4 Emerald Tint
COLOR_BORDER = "10B981"                   # #10B981 Emerald
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_p(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def create_document():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # COVER PAGE
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(30)
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
        
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_after = Pt(12)
    r_badge = p_badge.add_run("CV. EXADATA — LAPORAN KEMAJUAN PROGRES PERTENGAHAN")
    r_badge.font.name = 'Arial'
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_PRIMARY
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("LAPORAN ANTARA\nPENGEMBANGAN SOFTWARE APLIKASI PELAYANAN PERSAMPAHAN")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_DARK
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(35)
    r_sub = p_sub.add_run("Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) Terintegrasi\nKabupaten Banjarnegara, Jawa Tengah")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    
    div_table = doc.add_table(rows=1, cols=1)
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_background(div_cell, HEX_PRIMARY)
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    p_div = div_cell.paragraphs[0]
    p_div.paragraph_format.space_after = Pt(0)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(35)
    p_meta.paragraph_format.space_after = Pt(4)
    r_m1 = p_meta.add_run("Instansi Pemilik Pekerjaan: ")
    r_m1.font.bold = True
    p_meta.add_run("Dinas Perumahan, Kawasan Permukiman dan Lingkungan Hidup (DPPKPLH) Kabupaten Banjarnegara\n")
    
    p_meta2 = doc.add_paragraph()
    p_meta2.paragraph_format.space_after = Pt(4)
    r_m2 = p_meta2.add_run("Penyedia Jasa / Pelaksana: ")
    r_m2.font.bold = True
    p_meta2.add_run("CV. EXADATA\n")
    
    p_meta3 = doc.add_paragraph()
    p_meta3.paragraph_format.space_after = Pt(4)
    r_m3 = p_meta3.add_run("Capaian Progres Fisik: ")
    r_m3.font.bold = True
    p_meta3.add_run("65% (Minggu Ke-2 / Pertengahan Pelaksanaan)\n")
    
    p_meta4 = doc.add_paragraph()
    p_meta4.paragraph_format.space_after = Pt(0)
    r_m4 = p_meta4.add_run("Periode Pekerjaan: ")
    r_m4.font.bold = True
    p_meta4.add_run("1 – 30 Juli 2026 (Tahun Anggaran 2026)")

    doc.add_page_break()
    
    body_section = doc.sections[-1]
    header = body_section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("CV. EXADATA — Laporan Antara Pengembangan Software SIMPAH")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = COLOR_MUTED
    
    footer = body_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("DPPKPLH Kabupaten Banjarnegara | Dokumen Laporan Antara Pekerjaan")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = COLOR_MUTED

    # -------------------------------------------------------------
    # BAB 1 — PENDAHULUAN
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 1 — PENDAHULUAN")
    
    add_heading_2(doc, "1.1. Latar Belakang Laporan Antara")
    add_body_p(doc, "Pekerjaan Pengembangan Software Aplikasi Pelayanan Persampahan — Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) di Kabupaten Banjarnegara yang dilaksanakan oleh CV. EXADATA telah memasuki tahap pertengahan (Midterm Phase).")
    add_body_p(doc, "Laporan Antara ini disusun sebagai bentuk pertanggungjawaban progres pekerjaan teknis yang telah dicapai selama dua minggu pertama pelaksanaan proyek (periode 1 – 15 Juli 2026). Dokumen ini menyajikan hasil analisis kebutuhan terperinci, desain arsitektur software, skema database PostgreSQL Supabase, hasil perancangan UI/UX Progressive Web App (PWA), serta realisasi modul-modul utama yang telah selesai dikembangkan.")

    add_heading_2(doc, "1.2. Maksud dan Tujuan Laporan Antara")
    add_body_p(doc, "Maksud dari penyusunan Laporan Antara ini adalah untuk melaporkan realisasi fisik capaian pekerjaan pengembangan software kepada DPPKPLH Kabupaten Banjarnegara serta memverifikasi keselarasan fitur yang dibangun dengan kebutuhan operasional dinas.")
    add_body_p(doc, "Tujuan Laporan Antara meliputi:")
    add_bullet_p(doc, "Menyampaikan dokumentasi hasil analisis kebutuhan fungsional dan spesifikasi alur proses bisnis pengelolaan sampah.", "1. Pertanggungjawaban Teknis: ")
    add_bullet_p(doc, "Menyajikan desain arsitektur teknis PWA, struktur database Supabase PostgreSQL RLS, serta integrasi pemodelan Machine Learning Prophet dan AI Assistant Qwen.", "2. Transparansi Arsitektur Sistem: ")
    add_bullet_p(doc, "Melaporkan realisasi fisik pengembangan modul-modul software yang telah mencapai 65% dari total keseluruhan pekerjaan.", "3. Evaluasi Progres Fisik: ")
    add_bullet_p(doc, "Memperoleh masukan dan umpan balik dari tim teknis DPPKPLH Kabupaten Banjarnegara sebelum memasuki tahap akhir pengujian User Acceptance Testing (UAT) dan deployment.", "4. Review & Konfirmasi Dinas: ")

    add_heading_2(doc, "1.3. Ringkasan Capaian Progres Pertengahan (~65%)")
    add_body_p(doc, "Hingga minggu ke-2 pelaksanaan pekerjaan, CV. EXADATA telah menyelesaikan beberapa milestones utama:")
    
    add_callout_box(
        doc,
        "RINGKASAN CAPAIAN MILESTONES (PROGRES 65%)",
        [
            "Analisis Kebutuhan & Desain Arsitektur Sistem: Selesai 100% (Desain ERD Database, RBAC Matrix 4 Role, RLS Security Policies).",
            "Pengembangan Core PWA & Offline Engine: Selesai 100% (Service Worker, IndexedDB local storage, auto-sync initSync).",
            "Modul PWA Petugas Lapangan: Selesai 90% (Form Sampah Masuk Single/Batch, Input Pilah SIPSN 8 Kategori, Olah, Insidental, Fleet).",
            "Modul Validasi Koordinator Lapangan: Selesai 90% (Halaman Verifikasi Pending, Approve/Reject dengan Catatan Perbaikan).",
            "Modul Dashboard Eksekutif & ML Forecasting: Selesai 75% (4 KPI Cards, Grafik Tren & Integrasi API Machine Learning Meta Prophet).",
            "Modul Super Admin & Master Data: Selesai 70% (Manajemen Aduan & Tanggapan Dinas, Laporan Excel/SIPSN, Master Data Wilayah/Fasilitas/Armada/Kode Undangan)."
        ],
        bg_hex="F0FDF4",
        border_hex="10B981"
    )

    add_heading_2(doc, "1.4. Sistematika Penulisan Laporan Antara")
    add_body_p(doc, "Laporan Antara ini disusun dalam 6 BAB utama:")
    add_bullet_p(doc, "Pendahuluan, maksud, tujuan, ringkasan progres, dan sistematika.", "BAB 1 — PENDAHULUAN: ")
    add_bullet_p(doc, "Hasil analisis kebutuhan fungsional/non-fungsional dan workflow.", "BAB 2 — HASIL ANALISIS KEBUTUHAN SISTEM & SPESIFIKASI PROSES: ")
    add_bullet_p(doc, "Arsitektur PWA, ERD Database PostgreSQL, RLS Policies, ML Prophet, & AI Qwen.", "BAB 3 — DESAIN ARSITEKTUR SOFTWARE & SKEMA DATABASE: ")
    add_bullet_p(doc, "Detail hasil pembangunan modul PWA, Validasi, Eksekutif, Admin, & Portal Publik.", "BAB 4 — REALISASI PENGEMBANGAN FITUR & MODUL APLIKASI: ")
    add_bullet_p(doc, "Tabel perbandingan target vs realisasi, evaluasi Kurva S, dan kendala.", "BAB 5 — REKAPITULASI PROGRES PEKERJAAN & KURVA S: ")
    add_bullet_p(doc, "Rencana sisa pekerjaan minggu 3-4, UAT, training, & penutup.", "BAB 6 — RENCANA KERJA TAHAP AKHIR & PENUTUP: ")

    # -------------------------------------------------------------
    # BAB 2 — HASIL ANALISIS KEBUTUHAN SISTEM & SPESIFIKASI PROSES
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 2 — HASIL ANALISIS KEBUTUHAN SISTEM & SPESIFIKASI PROSES")
    
    add_heading_2(doc, "2.1. Analisis Kebutuhan Fungsional & Non-Fungsional")
    add_body_p(doc, "Berdasarkan hasil wawancara dan wawancara lapangan dengan tim teknis DPPKPLH Banjarnegara, dirumuskan spesifikasi kebutuhan software SIMPAH:")
    
    add_heading_3(doc, "A. Kebutuhan Fungsional (Functional Requirements):")
    add_bullet_p(doc, "Sistem harus menyediakan portal publik bebas akses untuk melihat statistik publik dan melacak nomor resi aduan.", "F-01: ")
    add_bullet_p(doc, "Sistem harus mendukung pendaftaran mandiri Warga dan pendaftaran Petugas dengan validasi Kode Undangan (Invitation Code).", "F-02: ")
    add_bullet_p(doc, "Sistem harus menyediakan form input PWA untuk Sampah Masuk (Single & Batch), Sampah Terpilah 8 Kategori SIPSN, Olah Sampah, Insidental, dan Armada Pengangkut.", "F-03: ")
    add_bullet_p(doc, "Sistem harus mewajibkan pengikatan koordinat GPS dan pengunggahan foto bukti pada setiap pengisian data.", "F-04: ")
    add_bullet_p(doc, "Sistem harus menyediakan halaman Validasi Data berjenjang bagi Koordinator Lapangan (Approve/Reject dengan catatan).", "F-05: ")
    add_bullet_p(doc, "Sistem harus menyediakan Dashboard Eksekutif dengan 4 KPI Cards, Grafik Tren Volume Sampah, dan Prediksi Machine Learning (Meta Prophet 7 hari ke depan).", "F-06: ")
    add_bullet_p(doc, "Sistem harus menyediakan modul Super Admin untuk mengelola Aduan & Tanggapan Dinas, Laporan Excel/SIPSN, 6 Tab Master Data, dan Manajemen MoU.", "F-07: ")

    add_heading_3(doc, "B. Kebutuhan Non-Fungsional (Non-Functional Requirements):")
    add_bullet_p(doc, "Waktu muat halaman (Page Load Time) di bawah 1.5 detik menggunakan bundler Vite modern.", "N-01 (Kinerja): ")
    add_bullet_p(doc, "Mendukung mode Offline-First di mana pengisian form tetap dapat dilakukan tanpa koneksi internet.", "N-02 (Aksesibilitas Offline): ")
    add_bullet_p(doc, "Keamanan data PostgreSQL dilindungi enkripsi HTTPS SSL dan Row Level Security (RLS).", "N-03 (Keamanan): ")
    add_bullet_p(doc, "Desain UI/UX Glassmorphic yang responsif dan fleksibel di smartphone, tablet, maupun desktop.", "N-04 (Usability): ")

    add_heading_2(doc, "2.2. Spesifikasi Alur Proses Bisnis & Data Flow")
    add_body_p(doc, "Alur proses bisnis pencatatan dan verifikasi data pada SIMPAH terbagi dalam 3 tahapan utama:")
    add_bullet_p(doc, "Petugas (Kader/Operator/Angkut) mengisi form di HP secara online atau offline. Data tersimpan dengan status 'Pending Verification' dan '_offlineSaved' jika tanpa sinyal.", "Tahap 1 (Penginputan Lapangan): ")
    add_bullet_p(doc, "Begitu perangkat online, data ter-sync ke Supabase. Koordinator Lapangan memeriksa data di menu Validasi. Data yang disetujui (Approved) resmi masuk ke statistik kabupaten, sedangkan data ditolak (Rejected) dikembalikan ke petugas.", "Tahap 2 (Verifikasi Berjenjang): ")
    add_bullet_p(doc, "Data yang terverifikasi secara otomatis dikalkulasi ke dalam Dashboard Eksekutif, Peta GIS, dan modul Ekspor Laporan SIPSN.", "Tahap 3 (Analitik & Pelaporan): ")

    # -------------------------------------------------------------
    # BAB 3 — DESAIN ARSITEKTUR SOFTWARE & SKEMA DATABASE
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 3 — DESAIN ARSITEKTUR SOFTWARE & SKEMA DATABASE")
    
    add_heading_2(doc, "3.1. Arsitektur Teknis Software PWA & Supabase")
    add_body_p(doc, "Software SIMPAH dirancang menggunakan arsitektur Decoupled Modern Web Application:")
    add_bullet_p(doc, "Menggunakan HTML5, Vanilla CSS, JavaScript ES6, dan Service Worker PWA.", "Frontend Tier: ")
    add_bullet_p(doc, "Supabase PostgreSQL BaaS (Backend-as-a-Service) yang menyediakan autentikasi terenkripsi dan API real-time.", "Database & Auth Tier: ")
    add_bullet_p(doc, "Python FastAPI + Meta Prophet ML model untuk layanan analitik prediktif.", "Machine Learning Tier: ")
    add_bullet_p(doc, "Alibaba Cloud MaaS API (Qwen LLM) untuk layanan AI Assistant SIMPAH Buddy.", "AI Integration Tier: ")

    add_heading_2(doc, "3.2. Desain Skema Database & Relasi Tabel (ERD)")
    add_body_p(doc, "Database PostgreSQL SIMPAH terdiri dari entitas tabel utama:")
    add_bullet_p(doc, "Menyimpan data profil pengguna, role (warga, petugas, eksekutif, admin), job_type, kecamatan, dan desa.", "1. Tabel `profiles`: ")
    add_bullet_p(doc, "Menyimpan seluruh transaksi sampah (Masuk, Pilah, Olah, Insidental) beserta berat kg, lokasi, foto, GPS, status_validasi, dan sync_action.", "2. Tabel `complaints` / `waste_inputs`: ")
    add_bullet_p(doc, "Menyimpan data referensi 28 Kecamatan dan 278 Desa/Kelurahan.", "3. Tabel `districts` & `villages`: ")
    add_bullet_p(doc, "Menyimpan data koordinat TPS, TPS3R, Bank Sampah, dan TPA Winong.", "4. Tabel `facilities` / `locations`: ")
    add_bullet_p(doc, "Menyimpan data plat nomor, jenis kendaraan, dan status armada.", "5. Tabel `fleets`: ")
    add_bullet_p(doc, "Menyimpan dokumen perjanjian kerja sama transporter dan status expired.", "6. Tabel `mou_contracts`: ")

    add_heading_2(doc, "3.3. Keamanan Data Row Level Security (RLS)")
    add_body_p(doc, "Penerapan RLS memastikan isolasi data secara ketat. Sebagai contoh, kebijakan `koordinator_select_policy` memastikan Koordinator hanya dapat membaca dan memvalidasi data yang berasal dari desa-desa di bawah wilayah kecamatannya (`kecamatan = auth.user.kecamatan`).")

    # -------------------------------------------------------------
    # BAB 4 — REALISASI PENGEMBANGAN FITUR & MODUL APLIKASI
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 4 — REALISASI PENGEMBANGAN FITUR & MODUL APLIKASI")
    
    add_heading_2(doc, "4.1. Realisasi Modul Portal Publik & Cek Resi Aduan")
    add_body_p(doc, "Modul Portal Publik (`#/portal/*`) telah selesai 100% dikembangkan. Fitur utama mencakup Landing Page resmi, Halaman Tentang SIMPAH, Galeri Edukasi Publik, Regulasi Kebijakan, Form Aduan Publik (Tanpa Login), serta Cek Status Resi Aduan dengan Progress Tracker 4 Tahap.")

    add_heading_2(doc, "4.2. Realisasi Modul PWA Petugas Lapangan & Offline Engine")
    add_body_p(doc, "Modul PWA Petugas telah mencapai progres 90%. Form Sampah Masuk (Single & Batch), Input Pilah SIPSN 8 Kategori, Olah Sampah, Insidental, Fleet Armada, dan Riwayat telah berfungsi lancar. Pengujian penginputan tanpa sinyal internet (Offline-First) dan auto-sync `initSync()` telah teruji sukses.")

    add_heading_2(doc, "4.3. Realisasi Modul Validasi Koordinator Lapangan")
    add_body_p(doc, "Modul Validasi (`#/dashboard/validasi`) telah mencapai progres 90%. Koordinator dapat meninjau antrean data pending, menyetujui data (Approve), atau menolak data (Reject) dengan menuliskan catatan perbaikan yang langsung muncul di HP petugas.")

    add_heading_2(doc, "4.4. Realisasi Modul Dashboard Eksekutif & ML Forecasting")
    add_body_p(doc, "Modul Eksekutif (`#/dashboard/eksekutif`) telah mencapai progres 75%. 4 KPI Cards, Grafik Tren Volume Sampah Harian, Integrasi API Machine Learning Meta Prophet (Garis Oranye Prediksi 7 Hari ke Depan), Donut Chart SIPSN, dan Peta GIS Spasial telah terintegrasi.")

    add_heading_2(doc, "4.5. Realisasi Modul Super Admin & Master Data")
    add_body_p(doc, "Modul Admin (`#/dashboard/*`) telah mencapai progres 70%. Manajemen Aduan & Tanggapan Dinas, Modul Ekspor Laporan Excel/SIPSN, 6 Tab Master Data (Wilayah, Lokasi, Armada, Kode Undangan, Penduduk, Pengguna), Manajemen MoU dengan Peringatan Expired, serta Widget AI SIMPAH Buddy telah aktif.")

    # -------------------------------------------------------------
    # BAB 5 — REKAPITULASI PROGRES PEKERJAAN & KURVA S
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 5 — REKAPITULASI PROGRES PEKERJAAN & KURVA S")
    
    add_heading_2(doc, "5.1. Tabel Perbandingan Target vs Realisasi Progres (Minggu 1-2)")
    add_body_p(doc, "Tabel di bawah ini menggambarkan perbandingan antara target rencana kerja dengan realisasi pencapaian fisik di lapangan hingga 15 Juli 2026:")

    headers_prog = ["No", "Komponen / Item Pekerjaan", "Bobot Target", "Realisasi", "Status"]
    data_prog = [
        ["1", "Analisis Kebutuhan & Desain Arsitektur Database", "15.0%", "15.0%", "Selesai 100%"],
        ["2", "Pengembangan Core PWA & Offline Engine", "15.0%", "15.0%", "Selesai 100%"],
        ["3", "Pengembangan Modul PWA Petugas Lapangan", "20.0%", "18.0%", "Selesai 90%"],
        ["4", "Pengembangan Modul Validasi Koordinator", "15.0%", "13.5%", "Selesai 90%"],
        ["5", "Pengembangan Dashboard Eksekutif & ML Prophet", "15.0%", "11.25%", "Selesai 75%"],
        ["6", "Pengembangan Super Admin & Master Data", "20.0%", "14.0%", "Selesai 70%"],
        ["", "TOTAL CAPAIAN PROGRES FISIK", "100.0%", "66.75%", "ON TRACK (65%+)"]
    ]

    table_p = doc.add_table(rows=len(data_prog) + 1, cols=5)
    table_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_p = table_p.rows[0].cells
    widths_p = [Inches(0.5), Inches(3.2), Inches(1.0), Inches(1.0), Inches(1.2)]
    for i, title in enumerate(headers_prog):
        hdr_cells_p[i].text = title
        hdr_cells_p[i].width = widths_p[i]
        set_cell_background(hdr_cells_p[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells_p[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_p[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data_prog):
        row_cells = table_p.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        is_total = row_idx == len(data_prog) - 1
        if is_total:
            bg_color = COLOR_LIGHT_BG
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_p[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if is_total:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_PRIMARY

    add_heading_2(doc, "5.2. Evaluasi Kendala Lapangan & Solusi Penanganan")
    add_body_p(doc, "Selama tahap pengembangan Minggu 1 dan 2, diidentifikasi kendala teknis penyesuaian kolom `response_text` pada tabel `complaints` Supabase. CV. EXADATA telah menyelesaikan perbaikan script migrasi database dan menyempurnakan logika sinkronisasi `updateComplaint` sehingga data tanggapan dinas dari Admin dapat tersimpan utuh dan dibaca langsung di HP Warga/Pelapor.")

    # -------------------------------------------------------------
    # BAB 6 — RENCANA KERJA TAHAP AKHIR & PENUTUP
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 6 — RENCANA KERJA TAHAP AKHIR & PENUTUP")
    
    add_heading_2(doc, "6.1. Rencana Pelaksanaan Sisa Pekerjaan (Minggu 3 – 4)")
    add_body_p(doc, "Pada sisa durasi pekerjaan (16 – 30 Juli 2026), CV. EXADATA fokus menyelesaikan tahapan finalisasi berikut:")
    add_bullet_p(doc, "Finalisasi penyempurnaan UI/UX modul Admin dan optimasi loading bundler PWA (16 – 20 Juli 2026).", "1. Penyempurnaan Sistem: ")
    add_bullet_p(doc, "Pelaksanaan simulasi UAT bersama tim teknis dan jajaran dinas DPPKPLH Banjarnegara (21 – 24 Juli 2026).", "2. Pengujian User Acceptance Testing (UAT): ")
    add_bullet_p(doc, "Pelaksanaan sesi pelatihan (training) operasional bagi Koordinator Lapangan, Operator TPS, dan Administrator (25 – 27 Juli 2026).", "3. Pelatihan & Pendampingan: ")
    add_bullet_p(doc, "Penyusunan Laporan Akhir, Buku Panduan Operasional (.docx & PDF), dan Serah Terima Pekerjaan (28 – 30 Juli 2026).", "4. Pelaporan Akhir & Serah Terima: ")

    add_heading_2(doc, "6.2. Kesimpulan Laporan Antara")
    add_body_p(doc, "Pekerjaan Pengembangan Software Aplikasi Pelayanan Persampahan (SIMPAH) oleh CV. EXADATA berjalan lancar, sesuai jadwal (on track), dan telah mencapai progres fisik 65%. Seluruh arsitektur core PWA, database Supabase PostgreSQL, ML Prophet, dan modul-modul utama telah berfungsi dengan baik dan siap memasuki tahap UAT serta serah terima akhir.")

    # Save document safely
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully created at primary path: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary file locked by Word. Document successfully saved at fallback path: {target_out}")

if __name__ == "__main__":
    create_document()
