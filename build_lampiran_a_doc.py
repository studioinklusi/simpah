import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Lampiran_A_Matriks_Hak_Akses_RBAC_dan_RLS.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Lampiran_A_Matriks_Hak_Akses_RBAC_dan_RLS_v1.docx"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # #059669 Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # #1F2937 Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # #6B7280 Muted Gray
COLOR_LIGHT_BG = "F0FDF4"                 # #F0FDF4 Emerald Tint
COLOR_BORDER = "10B981"                   # #10B981 Emerald
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"
HEX_HEADER_RLS = "1E293B"                # Dark Slate header for RLS table

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def build_document():
    doc = docx.Document()
    
    # Page setup - Margins 0.8 inch for wider tables
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # HEADER DOCUMENT
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("LAMPIRAN A: MATRIKS HAK AKSES RBAC & RLS SECURITY")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(18)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_s = p_sub.add_run("Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) Kabupaten Banjarnegara")
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(12)
    r_s.font.italic = True
    r_s.font.color.rgb = COLOR_MUTED

    # OVERVIEW CALLOUT
    add_callout_box(doc, "Ringkasan Arsitektur Keamanan Multi-Layer SIMPAH", [
        "Layer 1 — Authentication & Session: Otorisasi JWT via Supabase Auth dengan penjaminan kueri terproteksi SSL/TLS.",
        "Layer 2 — Backend Database Security (Row Level Security / RLS): Eksekusi kebijakan RLS native PostgreSQL pada Supabase yang membatasi hak baca & tulis data di tingkat baris database secara absolut.",
        "Layer 3 — Frontend Dynamic RBAC: Permission Matrix di antarmuka PWA/Dashboard berbasis tabel system_roles & role_permissions untuk menyembunyikan/menampilkan menu fitur secara dinamis."
    ])

    # BAGIAN 1: RBAC MATRIX
    add_heading_1(doc, "Bagian 1: Matriks Hak Akses RBAC (Role-Based Access Control)")
    add_body_p(doc, "Tabel berikut mendeskripsikan 7 tingkatan peran pengguna, cakupan data operasional, otorisasi tindakan CRUD (Create, Read, Update, Delete), serta fitur utama yang dapat diakses:")

    headers_rbac = ["Peran (Role)", "Sub-Peran / Job Type", "Scope Data Operasional", "CRUD", "Fitur Utama & Otoritas System"]
    data_rbac = [
        ["Super Admin", "admin", "Kabupaten (278 Desa & 28 Kec)", "C-R-U-D", "Full System Control, Master Data, MoU Transporter, Validasi Global, Audit Log, AI SIMPAH Buddy"],
        ["Eksekutif / Kadis", "eksekutif", "Kabupaten (Read-Only)", "- R - -", "Executive Dashboard (4 KPI Cards), Peta GIS Spasial, ML Prophet Timbulan, Donut SIPSN, Export Laporan"],
        ["Koordinator", "petugas (koordinator)", "Kecamatan Bertugas", "C-R-U-", "Validasi Data Lapangan (Approve/Reject + Catatan), Pengawasan Aduan Kecamatan, Monitoring Kader"],
        ["Operator TPS3R", "petugas (operator_tps)", "Fasilitas (location_id)", "C-R-U-", "Input Sampah Masuk TPS3R, Pemilahan 8 Kategori SIPSN, Pengolahan Kompos/Maggot, Catat Residu"],
        ["Petugas Angkut", "petugas (angkut)", "Armada (fleet_id)", "C-R-U-", "Input Timbangan Sampah Masuk, Ritase Pengangkutan ke TPA, Status & Pemeliharaan Armada"],
        ["Kader Lingkungan", "petugas (kader)", "Desa Bertugas (desa)", "C-R-U-", "Input Sampah Masuk RT/RW, Pilah & Olah Skala Warga, Lapor Kegiatan Insidental (Kerja Bakti)"],
        ["Warga Masyarakat", "warga", "Personal (auth.uid())", "C-R- -", "Kirim Aduan Tumpukan Sampah + Foto GPS, Melacak Status Resi Aduan, Peta GIS Fasum Kebersihan"]
    ]

    table_rbac = doc.add_table(rows=len(data_rbac) + 1, cols=5)
    table_rbac.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_rbac = [Inches(1.3), Inches(1.3), Inches(1.5), Inches(0.7), Inches(2.1)]
    
    # Format Header RBAC
    hdr_cells_r = table_rbac.rows[0].cells
    for i, title in enumerate(headers_rbac):
        hdr_cells_r[i].text = title
        hdr_cells_r[i].width = widths_rbac[i]
        set_cell_background(hdr_cells_r[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells_r[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_r[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 1, 3] else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Format Data RBAC
    for row_idx, row_data in enumerate(data_rbac):
        row_cells = table_rbac.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_rbac[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 1, 3] else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                if col_idx == 3:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # BAGIAN 2: RLS MATRIX TABLE
    add_heading_1(doc, "Bagian 2: Tabel Pengamanan Data Berbasis Row Level Security (RLS) PostgreSQL")
    add_body_p(doc, "Daftar kebijakan Row Level Security (RLS) terkonfigurasi pada PostgreSQL Supabase yang menjamin isolasi data secara mutlak:")

    headers_rls = ["Nama Tabel", "Nama Kebijakan RLS (Policy)", "Aksi", "Target Role", "Kriteria & Kondisi RLS Expression (USING / WITH CHECK)"]
    data_rls = [
        ["profiles", "public_read_profiles", "SELECT", "authenticated", "true (Pengguna terautentikasi dapat membaca profil umum pengguna lain)"],
        ["profiles", "update_own_profile", "UPDATE", "authenticated", "auth.uid() = id (Pengguna hanya bisa memperbarui profil pribadi)"],
        ["profiles", "admin_update_profiles", "UPDATE", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["waste_records", "petugas_insert", "INSERT", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role IN ('petugas', 'admin'))"],
        ["waste_records", "petugas_read_own", "SELECT", "authenticated", "created_by = auth.uid() OR village = (SELECT village FROM profiles WHERE id = auth.uid())"],
        ["waste_records", "koordinator_read_kecamatan", "SELECT", "authenticated", "kecamatan = (SELECT kecamatan FROM profiles WHERE id = auth.uid() AND job_type = 'koordinator')"],
        ["waste_records", "koordinator_update_kecamatan", "UPDATE", "authenticated", "kecamatan = (SELECT kecamatan FROM profiles WHERE id = auth.uid() AND job_type = 'koordinator')"],
        ["waste_records", "eksekutif_read_approved", "SELECT", "authenticated", "status = 'approved' AND EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'eksekutif')"],
        ["waste_records", "admin_full", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["sorted_waste", "read_sorted_waste", "SELECT", "authenticated", "true"],
        ["sorted_waste", "insert_sorted_waste", "INSERT", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role IN ('petugas', 'admin'))"],
        ["complaints", "public_insert", "INSERT", "anon, auth", "true (Formulir pengaduan warga publik terbuka tanpa kedala autentikasi)"],
        ["complaints", "user_read_own", "SELECT", "authenticated", "created_by = auth.uid() OR public_access = true"],
        ["complaints", "koordinator_admin_manage", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role IN ('admin', 'petugas'))"],
        ["incidental_events", "read_incidental_events", "SELECT", "authenticated", "true"],
        ["incidental_events", "insert_incidental_events", "INSERT", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role IN ('petugas', 'admin'))"],
        ["locations", "public_read_locations", "SELECT", "anon, auth", "true (Peta lokasi GIS TPS/TPS3R dapat diakses secara publik)"],
        ["locations", "admin_all_locations", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["fleet", "public_read_fleet", "SELECT", "anon, auth", "true"],
        ["fleet", "admin_all_fleet", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["mou & mou_fleet", "public_read_mou", "SELECT", "authenticated", "true"],
        ["mou & mou_fleet", "admin_all_mou", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["village_population", "Allow write admin only", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["public_facilities", "Allow write admin only", "ALL", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin')"],
        ["audit_log", "admin_select_audit_log", "SELECT", "authenticated", "EXISTS (SELECT 1 FROM profiles WHERE id = auth.uid() AND role = 'admin') (Jejak audit khusus Admin)"],
        ["audit_log", "insert_audit_log", "INSERT", "authenticated", "true (Pencatatan aktivitas otomatis oleh sistem)"]
    ]

    table_rls = doc.add_table(rows=len(data_rls) + 1, cols=5)
    table_rls.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_rls = [Inches(1.2), Inches(1.5), Inches(0.7), Inches(0.9), Inches(2.6)]

    # Format Header RLS
    hdr_cells_rls = table_rls.rows[0].cells
    for i, title in enumerate(headers_rls):
        hdr_cells_rls[i].text = title
        hdr_cells_rls[i].width = widths_rls[i]
        set_cell_background(hdr_cells_rls[i], HEX_HEADER_RLS)
        set_cell_margins(hdr_cells_rls[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_rls[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3] else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Format Data RLS
    for row_idx, row_data in enumerate(data_rls):
        row_cells = table_rls.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_rls[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [2, 3] else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8.5)
                if col_idx == 0:
                    run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # BAGIAN 3: SQL CODE REFERENCE
    add_heading_1(doc, "Bagian 3: Refrensi Skrip DDL PostgreSQL RLS (Supabase)")
    add_body_p(doc, "Berikut adalah potongan skrip DDL SQL yang diimplementasikan di Supabase SQL Editor untuk memperkuat sistem RLS:")

    sql_code = (
        "-- 1. Activate Row Level Security\n"
        "ALTER TABLE waste_records ENABLE ROW LEVEL SECURITY;\n"
        "ALTER TABLE profiles ENABLE ROW LEVEL SECURITY;\n"
        "ALTER TABLE complaints ENABLE ROW LEVEL SECURITY;\n\n"
        "-- 2. Trigger Protection: Prevent Role Escalation during Signup\n"
        "CREATE OR REPLACE FUNCTION handle_new_user()\n"
        "RETURNS TRIGGER AS $$\n"
        "BEGIN\n"
        "  INSERT INTO profiles (id, username, full_name, role)\n"
        "  VALUES (\n"
        "    new.id,\n"
        "    COALESCE(new.raw_user_meta_data->>'username', split_part(new.email, '@', 1)),\n"
        "    COALESCE(new.raw_user_meta_data->>'full_name', 'User Baru'),\n"
        "    'warga' -- Default role is strictly 'warga'\n"
        "  );\n"
        "  RETURN new;\n"
        "END;\n"
        "$$ LANGUAGE plpgsql SECURITY DEFINER;\n\n"
        "-- 3. Regional Scope RLS Policy for Koordinator Lapangan\n"
        "CREATE POLICY \"koordinator_read_kecamatan\" ON waste_records\n"
        "  FOR SELECT TO authenticated\n"
        "  USING (kecamatan = (SELECT kecamatan FROM profiles WHERE id = auth.uid() AND job_type = 'koordinator'));\n"
    )

    tbl_code = doc.add_table(rows=1, cols=1)
    tbl_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_code = tbl_code.cell(0, 0)
    c_code.width = Inches(6.5)
    set_cell_background(c_code, "1E293B") # Slate dark background
    set_cell_margins(c_code, top=120, bottom=120, left=150, right=150)
    p_c = c_code.paragraphs[0]
    p_c.paragraph_format.line_spacing = 1.1
    r_c = p_c.add_run(sql_code)
    r_c.font.name = 'Consolas'
    r_c.font.size = Pt(8.5)
    r_c.font.color.rgb = RGBColor(56, 189, 248) # Sky blue text

    # SAVE DOCUMENT
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully saved at: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary locked. Document saved at fallback: {target_out}")

if __name__ == "__main__":
    build_document()
