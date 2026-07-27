import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Lampiran_D_Glosarium_Istilah_Teknis_dan_Operasional.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Lampiran_D_Glosarium_Istilah_Teknis_dan_Operasional_v1.docx"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # Muted Gray
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"
HEX_HEADER_CAT = "1E293B"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📚 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def build_document():
    doc = docx.Document()
    
    # Page setup - Margins 0.75 inch
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # HEADER DOCUMENT
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("LAMPIRAN D: GLOSARIUM ISTILAH TEKNIS & OPERASIONAL")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(17)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_s = p_sub.add_run("Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) Kabupaten Banjarnegara")
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(11.5)
    r_s.font.italic = True
    r_s.font.color.rgb = COLOR_MUTED

    # OVERVIEW CALLOUT
    add_callout_box(doc, "Definisi Baku & Standarisasi Istilah SIMPAH", [
        "Tujuan Glosarium: Memberikan acuan pemahaman resmi terkait istilah rekayasa lunak, arsitektur data cloud, algoritma AI, serta operasional persampahan daerah.",
        "Kategori Istilah: Dibagi menjadi 3 kelompok utama: (1) Teknologi & Software Architecture, (2) Persampahan & Operasional Lapangan, (3) Artificial Intelligence & Analitik Data."
    ])

    # KATEGORI 1: TECH & SOFTWARE
    add_heading_1(doc, "Kategori 1: Istilah Teknologi & Arsitektur Software")
    add_body_p(doc, "Definisi teknis komponen arsitektur, keamanan, dan otentikasi sistem SIMPAH:")

    headers_tech = ["Istilah", "Kepanjangan / Istilah Asli", "Definisi Teknis & Peran Dalam Sistem SIMPAH"]
    data_tech = [
        ["PWA", "Progressive Web App", "Aplikasi web modern yang dapat dipasang di smartphone layaknya aplikasi native, mendukung kemampuan akses Offline Mode via Service Worker & IndexedDB."],
        ["RLS", "Row Level Security", "Fitur keamanan database PostgreSQL Supabase yang membatasi hak baca (SELECT) & tulis (INSERT/UPDATE/DELETE) data di tingkat baris secara absolut berdasarkan JWT role."],
        ["RBAC", "Role-Based Access Control", "Sistem kontrol akses berbasis peran (7 role) untuk mengatur visibilitas menu dan otorisasi tindakan pengguna."],
        ["IndexedDB", "Indexed Database API", "Database NoSQL lokal browser client untuk menyimpan transaksi data timbulan sampah secara sementara saat aplikasi bekerja tanpa koneksi internet (offline)."],
        ["JWT", "JSON Web Token", "Enkripsi token digital terstandar untuk otentikasi identitas pengguna dan pertukaran klaim role antara PWA frontend dan REST/Realtime API Supabase."],
        ["PWA Manifest", "Web App Manifest", "File JSON konfigurasi nama aplikasi, ikon, warna tema emerald, dan tampilan layar penuh (standalone) pada perangkat mobile."]
    ]

    tbl_tech = doc.add_table(rows=len(data_tech) + 1, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_tech = [Inches(1.2), Inches(1.8), Inches(3.8)]

    hdr_t = tbl_tech.rows[0].cells
    for i, title in enumerate(headers_tech):
        hdr_t[i].text = title
        hdr_t[i].width = widths_tech[i]
        set_cell_background(hdr_t[i], HEX_PRIMARY)
        set_cell_margins(hdr_t[i], top=90, bottom=90, left=100, right=100)
        p = hdr_t[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data_tech):
        row_cells = tbl_tech.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_tech[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8.5)
                if col_idx == 0:
                    run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # KATEGORI 2: OPERASIONAL PERSAMPAHAN
    add_heading_1(doc, "Kategori 2: Istilah Operasional Persampahan & Lapangan")
    add_body_p(doc, "Definisi baku istilah fasilitas, ritase, dan metode pengolahan sampah:")

    headers_ops = ["Istilah", "Kepanjangan / Istilah Asli", "Definisi Operasional & Peran Dalam Sistem SIMPAH"]
    data_ops = [
        ["SIPSN", "Sistem Informasi Pengelolaan Sampah Nasional", "Platform pelaporan resmi KLHK RI dengan 8 kategori sampah standar pelaporan nasional."],
        ["TPS3R", "Tempat Pengolahan Sampah 3R", "Fasilitas pengolahan sampah kawasan melayani pemilahan, pengomposan, pencacahan plastik, dan pembiakan maggot BSF."],
        ["Ritase", "Ritase Pengangkutan", "Jumlah frekuensi perjalanan bolak-balik armada angkut sampah dari TPS/Warga ke TPS3R atau TPA."],
        ["Residu", "Residu Sampah", "Sisa sampah non-recycle / tidak dapat diolah kembali yang wajib dibuang secara akhir ke TPA."],
        ["SLA", "Service Level Agreement", "Target batas waktu penanganan aduan masyarakat oleh dinas (misal: aduan selesai < 24 jam)."],
        ["Maggot BSF", "Bioconversion by Black Soldier Fly", "Pengolahan cepat sampah organik basah menggunakan larva lalat BSF menjadi pakan ternak berprotein tinggi."],
        ["Eco-Enzyme", "Fermentation Eco-Enzyme", "Cairan hasil fermentasi sampah organik basah (kulit buah/sayur), gula merah, dan air untuk pembersih/disinfektan alami."],
        ["RDF", "Refuse Derived Fuel", "Bahan bakar alternatif berbahan cacahan sampah kering berkalori tinggi pengganti batu bara di industri semen/PLTSa."]
    ]

    tbl_ops = doc.add_table(rows=len(data_ops) + 1, cols=3)
    tbl_ops.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_ops = [Inches(1.2), Inches(1.8), Inches(3.8)]

    hdr_o = tbl_ops.rows[0].cells
    for i, title in enumerate(headers_ops):
        hdr_o[i].text = title
        hdr_o[i].width = widths_ops[i]
        set_cell_background(hdr_o[i], HEX_HEADER_CAT)
        set_cell_margins(hdr_o[i], top=90, bottom=90, left=100, right=100)
        p = hdr_o[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data_ops):
        row_cells = tbl_ops.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_ops[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8.5)
                if col_idx == 0:
                    run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # KATEGORI 3: AI & ANALYTICS
    add_heading_1(doc, "Kategori 3: Istilah Artificial Intelligence & Analitik Data")
    add_body_p(doc, "Definisi istilah kecerdasan buatan, permodelan bahasa, dan forecasting timbulan:")

    headers_ai = ["Istilah", "Kepanjangan / Istilah Asli", "Definisi Teknis & Peran Dalam Sistem SIMPAH"]
    data_ai = [
        ["DSS", "Decision Support System", "Sistem pendukung keputusan terintegrasi (grafik analitik + GIS spasial + forecasting) untuk rekomendasi kebijakan pimpinan."],
        ["LLM", "Large Language Model", "Model AI berbasis deep learning yang dilatih pada miliaran data teks untuk memahami & merespons percakapan bahasa alami."],
        ["Qwen", "Qwen LLM (Alibaba Cloud MaaS)", "Model LLM canggih Alibaba Cloud yang diintegrasikan pada widget SIMPAH Buddy untuk asisten interaktif persampahan."],
        ["Meta Prophet", "Meta Prophet Forecasting", "Algoritma machine learning time-series kembangan Meta pada microservice Python FastAPI untuk memprediksi proyeksi timbulan harian."]
    ]

    tbl_ai = doc.add_table(rows=len(data_ai) + 1, cols=3)
    tbl_ai.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_ai = [Inches(1.2), Inches(1.8), Inches(3.8)]

    hdr_a = tbl_ai.rows[0].cells
    for i, title in enumerate(headers_ai):
        hdr_a[i].text = title
        hdr_a[i].width = widths_ai[i]
        set_cell_background(hdr_a[i], HEX_PRIMARY)
        set_cell_margins(hdr_a[i], top=90, bottom=90, left=100, right=100)
        p = hdr_a[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data_ai):
        row_cells = tbl_ai.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_ai[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8.5)
                if col_idx == 0:
                    run.font.bold = True

    # SAVE DOCUMENT
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully saved at: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary locked. Document saved at fallback: {target_out}")

if __name__ == "__main__":
    build_document()
