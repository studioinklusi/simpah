import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Panduan_Operasional_SIMPAH.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Panduan_Operasional_SIMPAH_v1.docx"
LOGO_PATH = r"u:\Project\simpah-rilis v1\logo.png"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # #059669 Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # #1F2937 Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # #6B7280 Muted Gray
COLOR_LIGHT_BG = "F0FDF4"                 # #F0FDF4 Emerald Tint
COLOR_BORDER = "10B981"                   # #10B981 Emerald
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_p(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p

def add_screenshot_placeholder(doc, code, title, description, size_hint="Rekomendasi ukuran: Layar Penuh (16:9) atau Mobile Frame (9:16)"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    
    set_cell_background(cell, COLOR_LIGHT_BG)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': COLOR_BORDER, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_code = p.add_run(f"📷 PLACEHOLDER SCREENSHOT [{code}]\n")
    run_code.font.name = 'Arial'
    run_code.font.size = Pt(10)
    run_code.font.bold = True
    run_code.font.color.rgb = COLOR_PRIMARY
    
    run_title = p.add_run(f"Judul Tampilan: {title}\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_DARK
    
    run_desc = p.add_run(f"Keterangan Gambar: {description}\n")
    run_desc.font.name = 'Calibri'
    run_desc.font.size = Pt(9.5)
    run_desc.font.color.rgb = COLOR_DARK
    
    run_hint = p.add_run(f"📌 {size_hint}")
    run_hint.font.name = 'Calibri'
    run_hint.font.size = Pt(9)
    run_hint.font.italic = True
    run_hint.font.color.rgb = COLOR_MUTED
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def create_document():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # COVER PAGE
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(30)
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
        
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_after = Pt(12)
    r_badge = p_badge.add_run("STUDIO INKLUSI — PANDUAN OPERASIONAL")
    r_badge.font.name = 'Arial'
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_PRIMARY
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("SIMPAH\nSistem Informasi Manajemen Pengelolaan Sampah")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(26)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_DARK
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(40)
    r_sub = p_sub.add_run("Buku Panduan Penggunaan & Operasional Sistem Terintegrasi untuk Warga, Petugas Lapangan, Koordinator, Eksekutif, dan Administrator")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    
    div_table = doc.add_table(rows=1, cols=1)
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_background(div_cell, HEX_PRIMARY)
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    p_div = div_cell.paragraphs[0]
    p_div.paragraph_format.space_after = Pt(0)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(40)
    p_meta.paragraph_format.space_after = Pt(4)
    r_m1 = p_meta.add_run("Pengembang Sistem: ")
    r_m1.font.bold = True
    p_meta.add_run("STUDIO INKLUSI\n")
    
    p_meta2 = doc.add_paragraph()
    p_meta2.paragraph_format.space_after = Pt(4)
    r_m2 = p_meta2.add_run("Versi Dokumen: ")
    r_m2.font.bold = True
    p_meta2.add_run("1.0 (Rilis Resmi)\n")
    
    p_meta3 = doc.add_paragraph()
    p_meta3.paragraph_format.space_after = Pt(4)
    r_m3 = p_meta3.add_run("Tanggal Update: ")
    r_m3.font.bold = True
    p_meta3.add_run("Juli 2026\n")
    
    p_meta4 = doc.add_paragraph()
    p_meta4.paragraph_format.space_after = Pt(0)
    r_m4 = p_meta4.add_run("Wilayah Implementasi: ")
    r_m4.font.bold = True
    p_meta4.add_run("Kabupaten Banjarnegara, Jawa Tengah")

    doc.add_page_break()
    
    body_section = doc.sections[-1]
    header = body_section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("STUDIO INKLUSI — Panduan Operasional Aplikasi SIMPAH")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = COLOR_MUTED
    
    footer = body_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("Dokumen Resmi SIMPAH v1.0 | Halaman Ref Dokumen")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = COLOR_MUTED

    # BAB 1
    add_heading_1(doc, "BAB 1 — PENDAHULUAN")
    add_heading_2(doc, "1.1. Apa itu SIMPAH?")
    add_body_p(doc, "SIMPAH (Sistem Informasi Manajemen Pengelolaan Sampah) adalah platform digital terintegrasi berbasis Progressive Web App (PWA) yang dirancang khusus untuk mendigitalisasi dan memonitor seluruh rantai pasok pengelolaan sampah secara real-time. Sistem ini mencakup pencatatan dari tingkat sumber (warga dan TPS), pengangkutan oleh armada, validasi lapangan, hingga pengolahan di fasilitas akhir (TPS3R, Bank Sampah, dan TPA).")
    add_body_p(doc, "SIMPAH dilengkapi dengan kecerdasan buatan AI Assistant (SIMPAH Buddy) dan pemodelan prediktif Machine Learning (Meta Prophet) untuk mendukung pengambilan keputusan berbasis data yang akurat, transparan, serta selaras dengan standar Pelaporan Informasi Pengelolaan Sampah Nasional (SIPSN).")

    add_heading_2(doc, "1.2. Tujuan Aplikasi")
    add_body_p(doc, "Implementasi SIMPAH bertujuan untuk mencapai beberapa sasaran strategis berikut:")
    add_bullet_p(doc, "Memastikan seluruh pencatatan berat dan kategori sampah terverifikasi secara berjenjang tanpa risiko manipulasi data.", "1. Akurasi & Validitas Data: ")
    add_bullet_p(doc, "Menyediakan aplikasi berbasis PWA yang mendukung penginputan data secara offline di lokasi minim sinyal dan otomatis menyinkronkan data saat terhubung kembali ke internet.", "2. Aksesibilitas & Offline-First: ")
    add_bullet_p(doc, "Menyediakan ringkasan metrik utama (KPI), peta titik sampah GIS, serta laporan terstruktur yang siap diunduh untuk kebutuhan dinas.", "3. Transparansi & Dashboard Eksekutif: ")
    add_bullet_p(doc, "Memanfaatkan algoritma Machine Learning untuk memprediksi tren volume sampah 7 hari ke depan guna mengoptimalkan alokasi armada pengangkut.", "4. Analitik Prediktif (ML Forecasting): ")
    add_bullet_p(doc, "Memudahkan masyarakat dalam menyampaikan aduan penumpukan sampah atau pembuangan liar berbasis foto dan koordinat GPS.", "5. Pemberdayaan Masyarakat & Pelayanan Publik: ")

    add_heading_2(doc, "1.3. Arsitektur Sistem")
    add_body_p(doc, "SIMPAH dibangun di atas arsitektur teknologi modern yang andal, responsif, dan aman:")
    add_bullet_p(doc, "Menggunakan teknologi web modern (HTML5, JavaScript ES6, Vanilla CSS) yang di-build sebagai PWA ultra-cepat. Didesain dengan prinsip Mobile-First dan Glassmorphism UI yang futuristik.", "Frontend PWA: ")
    add_bullet_p(doc, "Database PostgreSQL terenkripsi dengan otentikasi aman dan sistem keamanan tingkat baris data (Row Level Security / RLS).", "Database & Authentication (Supabase): ")
    add_bullet_p(doc, "Penyimpanan data lokal di HP/perangkat menggunakan IndexedDB dan Service Worker, memungkinkan pengisian form tetap berjalan lancar saat tidak ada sinyal.", "Offline-First Engine: ")
    add_bullet_p(doc, "Layanan backend Python (FastAPI + Meta Prophet) untuk menghitung prediksi data volume sampah harian.", "Machine Learning Backend: ")
    add_bullet_p(doc, "Model bahasa besar (LLM) Qwen dari Alibaba Cloud yang siap memberikan edukasi, ringkasan laporan, dan panduan penggunaan langsung di dalam aplikasi.", "AI Chatbot (SIMPAH Buddy): ")

    add_heading_2(doc, "1.4. Persyaratan Perangkat & Browser")
    add_body_p(doc, "Untuk menggunakan SIMPAH secara optimal, perangkat yang digunakan hendaknya memenuhi spesifikasi minimum berikut:")
    add_bullet_p(doc, "Google Chrome (versi 90+), Safari (iOS 13+), atau Microsoft Edge. Browser harus mendukung fitur Service Worker dan Geolocation.", "Browser Rekomendasi: ")
    add_bullet_p(doc, "Android 8.0 (Oreo) ke atas atau iOS 13 ke atas, RAM minimal 2 GB, serta memiliki GPS/Fitur Lokasi aktif.", "Perangkat Mobile (HP): ")
    add_bullet_p(doc, "Layar minimal 1280x720 piksel (rekomendasi Full HD 1920x1080) untuk melihat Dashboard Eksekutif dan Peta GIS secara maksimal.", "Perangkat Komputer/Laptop: ")
    add_bullet_p(doc, "Koneksi internet (3G/4G/WiFi) diperlukan saat login awal dan proses sinkronisasi data.", "Koneksi Internet: ")

    add_heading_2(doc, "1.5. Cara Mengakses Aplikasi")
    add_body_p(doc, "Aplikasi SIMPAH dapat diakses melalui alamat URL resmi berikut:")
    add_bullet_p(doc, "https://www.simpah.id", "Alamat Portal Publik & Login: ")
    add_body_p(doc, "Pengguna dapat mengklik tombol 'Masuk Sistem' di pojok kanan atas portal untuk menuju halaman Login pengguna terdaftar.")

    add_screenshot_placeholder(
        doc, 
        "SS-01", 
        "Halaman Utama Portal Publik SIMPAH", 
        "Tampilan beranda portal publik (www.simpah.id) memperlihatkan banner utama 'Smart Waste Intelligence', tombol Masuk Sistem, serta navigasi menu portal publik.",
        "Rekomendasi ukuran: Desktop Full Width (16:9)"
    )

    add_heading_2(doc, "1.6. Cara Install PWA di Smartphone")
    add_body_p(doc, "Sebagai aplikasi berbasis PWA (Progressive Web App), SIMPAH dapat diinstal langsung ke layar utama (Home Screen) smartphone Android maupun iPhone tanpa harus mengunduh file besar dari Google Play Store atau Apple App Store.")
    
    add_heading_3(doc, "Langkah Install di Android (Google Chrome):")
    add_bullet_p(doc, "Buka browser Google Chrome di HP Android Anda dan kunjungi https://www.simpah.id.")
    add_bullet_p(doc, "Tunggu beberapa detik hingga muncul banner notifikasi di bawah layar: 'Tambahkan SIMPAH ke Layar Utama'.")
    add_bullet_p(doc, "Jika banner tidak muncul, ketuk titik tiga (⋮) di pojok kanan atas Chrome, lalu pilih menu 'Install aplikasi' atau 'Tambahkan ke Layar Utama'.")
    add_bullet_p(doc, "Ketuk tombol 'Install' / 'Tambah' pada pop-up konfirmasi.")
    add_bullet_p(doc, "Icon aplikasi SIMPAH akan muncul di layar utama HP Anda dan dapat dibuka layaknya aplikasi native.")

    add_heading_3(doc, "Langkah Install di iPhone / iPad (Safari):")
    add_bullet_p(doc, "Buka browser Safari di iPhone Anda dan kunjungi https://www.simpah.id.")
    add_bullet_p(doc, "Ketuk ikon 'Share' (kotak dengan panah ke atas) di bagian bawah layar Safari.")
    add_bullet_p(doc, "Gulir ke bawah dan ketuk opsi 'Add to Home Screen' (Tambahkan ke Layar Utama).")
    add_bullet_p(doc, "Ketuk tombol 'Add' di pojok kanan atas.")

    add_screenshot_placeholder(
        doc, 
        "SS-02", 
        "Prompt / Banner Install PWA di Smartphone", 
        "Tampilan layar smartphone yang memperlihatkan banner atau pop-up tombol 'Tambahkan SIMPAH ke Layar Utama' (Add to Home Screen).",
        "Rekomendasi ukuran: Smartphone Screen (9:16)"
    )

    add_screenshot_placeholder(
        doc, 
        "SS-03", 
        "Icon SIMPAH di Layar Utama (Home Screen) HP", 
        "Tampilan layar beranda (home screen) smartphone yang menunjukkan icon aplikasi SIMPAH yang berhasil terpasang.",
        "Rekomendasi ukuran: Smartphone Screen (9:16)"
    )

    # BAB 2
    add_heading_1(doc, "BAB 2 — SISTEM PERAN & HAK AKSES")
    add_heading_2(doc, "2.1. Daftar Peran Pengguna (Role System)")
    add_body_p(doc, "SIMPAH menerapkan sistem kontrol akses berbasis peran (Role-Based Access Control / RBAC) yang membagi pengguna ke dalam 4 tingkatan peran utama dan 4 sub-peran petugas lapangan:")
    add_bullet_p(doc, "Masyarakat umum yang terdaftar. Memiliki wewenang untuk mengirimkan laporan pengaduan sampah, melihat riwayat aduan sendiri, dan mengakses fitur edukasi.", "1. Peran Warga: ")
    add_bullet_p(doc, "Petugas teknis di lapangan yang menangani pencatatan, pengangkutan, pengolahan, dan verifikasi data. Terdiri dari 4 sub-peran (Job Types):", "2. Peran Petugas Lapangan: ")
    add_bullet_p(doc, "Petugas penggerak di tingkat RT/RW/Desa yang mencatat timbulan sampah rumah tangga dan pengolahan warga.", "   • Kader Lingkungan (kader): ")
    add_bullet_p(doc, "Petugas di fasilitas TPS3R, Rumah Kompos, atau Bank Sampah yang menginput sampah masuk, pemilahan, dan residu fasilitas.", "   • Operator TPS3R / Bank Sampah (operator_tps): ")
    add_bullet_p(doc, "Sopir atau pengangkut armada truk/tossa sampah yang mencatat ritase pengangkutan dan status armada.", "   • Petugas Pengangkut / Armada (angkut): ")
    add_bullet_p(doc, "Pengawas tingkat kecamatan yang bertugas menginventarisasi dan memvalidasi (menyetujui/menolak) seluruh data inputan petugas sebelum masuk ke dashboard resmi.", "   • Koordinator Lapangan (koordinator): ")
    add_bullet_p(doc, "Pimpinan daerah / Kepala Dinas Lingkungan Hidup. Memiliki hak akses khusus untuk melihat Dashboard Eksekutif, statistik KPI, forecasting ML, dan peta GIS.", "3. Peran Eksekutif: ")
    add_bullet_p(doc, "Pengelola utama sistem yang memiliki wewenang penuh untuk mengelola master data wilayah, pengguna, armada, MoU, artikel edukasi, validasi data, serta mengunduh laporan SIPSN.", "4. Peran Admin: ")

    add_heading_2(doc, "2.2. Matriks Hak Akses (Permission Matrix)")
    add_body_p(doc, "Tabel di bawah ini menjelaskan pembagian hak akses fitur pada setiap peran pengguna di dalam aplikasi SIMPAH:")

    headers = ["Modul / Fitur Aplikasi", "Warga", "Petugas", "Koordinator", "Eksekutif", "Admin"]
    data_perm = [
        ["Portal Publik & Cek Resi Aduan", "✓", "✓", "✓", "✓", "✓"],
        ["Buat Aduan Masyarakat & Foto GPS", "✓", "✓", "✓", "✓", "✓"],
        ["Lihat Aduan Saya (Milik Sendiri)", "✓", "✓", "✓", "✓", "✓"],
        ["Input Sampah Masuk / Pilah / Olah", "—", "✓*", "—", "—", "✓"],
        ["Input Sampah Insidental & Fleet", "—", "✓*", "—", "—", "✓"],
        ["Validasi Data Lapangan (Approve/Reject)", "—", "—", "✓", "—", "✓"],
        ["Manajemen Status Aduan & Tanggapan", "—", "—", "—", "✓", "✓"],
        ["Dashboard Eksekutif & ML Forecasting", "—", "—", "—", "✓", "✓"],
        ["Peta GIS Distribusi Titik Sampah", "✓", "✓", "✓", "✓", "✓"],
        ["Manajemen Master Data & Kode Undangan", "—", "—", "—", "—", "✓"],
        ["Manajemen MoU Transporter & Fasilitas", "—", "—", "—", "—", "✓"],
        ["Ekspor Laporan (Excel / SIPSN)", "—", "—", "—", "—", "✓"],
        ["SIMPAH Buddy (AI Chat Assistant)", "—", "✓", "✓", "✓", "✓"]
    ]
    
    table_p = doc.add_table(rows=len(data_perm) + 1, cols=6)
    table_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_p.rows[0].cells
    widths = [Inches(2.5), Inches(0.8), Inches(0.8), Inches(0.9), Inches(0.8), Inches(0.8)]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data_perm):
        row_cells = table_p.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if cell_value == "✓":
                    run.font.bold = True
                    run.font.color.rgb = COLOR_PRIMARY
                elif cell_value == "—":
                    run.font.color.rgb = COLOR_MUTED

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(4)
    p_note.paragraph_format.space_after = Pt(12)
    r_n = p_note.add_run("* Catatan: Input data lapangan disesuaikan dengan Jenis Tugas (Job Type) masing-masing petugas.")
    r_n.font.name = 'Calibri'
    r_n.font.size = Pt(9)
    r_n.font.italic = True
    r_n.font.color.rgb = COLOR_MUTED

    add_heading_2(doc, "2.3. Cara Registrasi Akun Baru")
    add_body_p(doc, "Pengguna baru dapat mendaftarkan diri secara mandiri melalui form registrasi di aplikasi SIMPAH. Prosedur pendaftaran dibedakan berdasarkan peran yang dipilih:")
    add_heading_3(doc, "A. Pendaftaran Akun Warga:")
    add_bullet_p(doc, "Buka aplikasi SIMPAH dan klik tombol 'Masuk Sistem' lalu pilih 'Daftar Sekarang' (atau buka URL https://www.simpah.id/#/register).")
    add_bullet_p(doc, "Isi Nama Lengkap, Nomor Telepon/WhatsApp, Username, Email, dan Password.")
    add_bullet_p(doc, "Pilih Peran: 'Warga / Masyarakat'.")
    add_bullet_p(doc, "Pilih Wilayah Domisili (Kecamatan dan Desa/Kelurahan).")
    add_bullet_p(doc, "Ketuk tombol 'Daftar Akun Baru'. Akun Warga langsung aktif dan siap digunakan.")

    add_heading_3(doc, "B. Pendaftaran Akun Petugas Lapangan:")
    add_bullet_p(doc, "Pilih Peran: 'Petugas Lapangan'.")
    add_bullet_p(doc, "Pilih Jenis Tugas (Job Type): Kader Lingkungan / Operator TPS3R / Petugas Angkut / Koordinator Lapangan.")
    add_bullet_p(doc, "Masukkan Kode Undangan Resmi (Invitation Code) yang diberikan oleh Administrator Dinas. Kode ini wajib diisi sebagai validasi keamanan otorisasi petugas.")
    add_bullet_p(doc, "Lengkapi data domisili/fasilitas tempat bertugas, lalu klik 'Daftar Akun Baru'.")

    add_screenshot_placeholder(
        doc, 
        "SS-05", 
        "Formulir Registrasi Akun Baru SIMPAH", 
        "Tampilan form registrasi memperlihatkan field Nama Lengkap, Username, Password, pilihan Peran (Role), pilihan Jenis Tugas, dan input Kode Undangan untuk Petugas.",
        "Rekomendasi ukuran: Mobile / Laptop Screen"
    )

    add_heading_2(doc, "2.4. Cara Login Ke Aplikasi")
    add_body_p(doc, "Langkah-langkah untuk masuk ke akun SIMPAH yang telah terdaftar:")
    add_bullet_p(doc, "Buka aplikasi SIMPAH di browser HP/Laptop Anda.")
    add_bullet_p(doc, "Masukkan Username (contoh: kader1) atau Email lengkap (contoh: kader1@simpah.dev). Catatan: Sistem secara otomatis akan menambahkan domain @simpah.dev jika Anda hanya memasukkan username saja.", "1. Input Username/Email: ")
    add_bullet_p(doc, "Masukkan password akun Anda dengan benar.", "2. Input Password: ")
    add_bullet_p(doc, "Ketuk tombol 'Masuk Sistem'. Sistem akan mengarahkan Anda ke Halaman Beranda sesuai peran masing-masing secara otomatis.", "3. Masuk: ")

    add_screenshot_placeholder(
        doc, 
        "SS-04", 
        "Halaman Login SIMPAH", 
        "Tampilan form login aplikasi SIMPAH yang bersih dengan input Username/Email, Password, tombol Masuk Sistem, serta link Lupa Password.",
        "Rekomendasi ukuran: Laptop / Tablet Screen (16:9)"
    )

    add_heading_2(doc, "2.5. Lupa Password & Reset Password")
    add_body_p(doc, "Jika Anda lupa kata sandi untuk masuk ke sistem, ikuti langkah pemulihan berikut:")
    add_bullet_p(doc, "Pada halaman Login, ketuk tautan 'Lupa Password?'.")
    add_bullet_p(doc, "Masukkan alamat Email terdaftar yang Anda gunakan saat registrasi.")
    add_bullet_p(doc, "Ketuk tombol 'Kirim Instruksi Reset'. Sistem akan mengirimkan tautan pemulihan sandi ke email Anda.")
    add_bullet_p(doc, "Buka email Anda, klik tautan yang diterima, lalu masukkan kata sandi baru Anda.")

    add_screenshot_placeholder(
        doc, 
        "SS-06", 
        "Halaman Permintaan Lupa Password", 
        "Tampilan form Lupa Password tempat pengguna memasukkan alamat email terdaftar untuk menerima link reset kata sandi.",
        "Rekomendasi ukuran: Laptop / Tablet Screen"
    )

    add_screenshot_placeholder(
        doc, 
        "SS-07", 
        "Konfirmasi Kirim Tautan Reset Password", 
        "Tampilan notifikasi berhasil dikirimnya instruksi reset kata sandi ke email pengguna.",
        "Rekomendasi ukuran: Mobile / Tablet Screen"
    )

    # BAB 3 - BAB 9
    add_heading_1(doc, "BAB 3 — PANDUAN UNTUK PERAN: WARGA")
    add_heading_2(doc, "3.1. Beranda Warga")
    add_body_p(doc, "Beranda PWA Warga menyajikan ringkasan aduan, menu cepat pembuatan aduan baru, dan riwayat aduan terakhir.")
    add_screenshot_placeholder(doc, "SS-08", "Beranda PWA Tampilan Peran Warga", "Tampilan beranda PWA Warga.", "Smartphone Screen (9:16)")

    add_heading_2(doc, "3.2. Membuat Aduan Baru")
    add_body_p(doc, "Alur 9 langkah pengaduan sampah berbasis GPS, foto, pilihan 7 kategori masalah, mode anonim, dan nomor resi.")
    add_screenshot_placeholder(doc, "SS-09", "Formulir Buat Aduan Baru (Warga)", "Form aduan baru warga.", "Smartphone Screen (9:16)")

    add_heading_2(doc, "3.3. Melihat Status & Riwayat Aduan Saya")
    add_body_p(doc, "Memantau status penanganan aduan (Baru, Diproses, Ditindaklanjuti, Selesai, Ditolak) serta Tanggapan Dinas.")
    add_screenshot_placeholder(doc, "SS-10", "Detail Aduan Warga & Tanggapan Dinas", "Modal detail aduan dan tanggapan dinas.", "Mobile / Tablet Screen")

    add_heading_2(doc, "3.4. Melacak Aduan via Portal Publik")
    add_body_p(doc, "Pelacakan nomor resi aduan di portal publik tanpa login dilengkapi Progress Tracker 4 Tahap.")
    add_screenshot_placeholder(doc, "SS-11", "Halaman Pelacakan Aduan Portal Publik", "Form input resi portal publik.", "Desktop / Tablet Screen")
    add_screenshot_placeholder(doc, "SS-12", "Hasil Pelacakan Aduan & Progress Bar Tracker", "Hasil pelacakan aduan resi.", "Tablet / Desktop Screen")

    add_heading_2(doc, "3.5. Menu Edukasi & Informasi Sampah")
    add_body_p(doc, "Membaca artikel panduan pemilahan sampah organik vs anorganik dan komposting.")

    add_heading_1(doc, "BAB 4 — PANDUAN UNTUK PERAN: KADER LINGKUNGAN")
    add_heading_2(doc, "4.1. Beranda Kader Lingkungan")
    add_body_p(doc, "Wewenang desa, 4 kartu statistik operasional, dan menu cepat pencatatan.")
    add_screenshot_placeholder(doc, "SS-13", "Beranda PWA Tampilan Kader Lingkungan", "Beranda PWA Kader.", "Smartphone Screen (9:16)")

    add_heading_2(doc, "4.2. Input Sampah Masuk (Campur)")
    add_body_p(doc, "Input Mode Single (1 hari) & Mode Batch (beberapa hari), berat (kg), lokasi, foto, GPS.")
    add_screenshot_placeholder(doc, "SS-14", "Formulir Input Sampah Masuk (Kader)", "Form input sampah masuk.", "Smartphone Screen (9:16)")
    add_screenshot_placeholder(doc, "SS-15", "Formulir Input Sampah Masuk Terisi Lengkap", "Form terisi foto & GPS.", "Smartphone Screen (9:16)")

    add_heading_2(doc, "4.3. Input Sampah Terpilah (Standar SIPSN)")
    add_body_p(doc, "Pengisian 8 kategori standar SIPSN dengan kalkulasi total otomatis.")
    add_screenshot_placeholder(doc, "SS-16", "Formulir Input Sampah Terpilah Kategori SIPSN", "Form input sampah terpilah SIPSN.", "Mobile / Tablet Screen")

    add_heading_2(doc, "4.4. Input Olah Sampah & Kejadian Insidental")
    add_body_p(doc, "Pencatatan pengolahan sampah dan pencatatan kegiatan gotong royong/insidental.")
    add_screenshot_placeholder(doc, "SS-17", "Formulir Input Pengolahan Sampah", "Form input olah sampah.", "Mobile Screen (9:16)")
    add_screenshot_placeholder(doc, "SS-18", "Formulir Pencatatan Kejadian Insidental", "Form pencatatan insidental.", "Mobile Screen (9:16)")

    add_heading_2(doc, "4.5. Riwayat Pencatatan & Mode Offline Auto-Sync")
    add_body_p(doc, "Daftar riwayat verifikasi dan mekanisme sinkronisasi otomatis saat terhubung kembali ke internet.")
    add_screenshot_placeholder(doc, "SS-19", "Halaman Riwayat Pencatatan Kader", "Halaman riwayat pencatatan.", "Mobile Screen (9:16)")
    add_screenshot_placeholder(doc, "SS-20", "Indikator Status Mode Offline di PWA", "Header mode offline.", "Smartphone Header View")
    add_screenshot_placeholder(doc, "SS-21", "Notifikasi Sinkronisasi Data Berhasil", "Toast notifikasi sync berhasil.", "Mobile Screen Notification")

    add_heading_1(doc, "BAB 5 — PANDUAN UNTUK PERAN: OPERATOR TPS3R / BANK SAMPAH")
    add_heading_2(doc, "5.1. Beranda Operator TPS3R")
    add_body_p(doc, "Lencana fasilitas terikat dan statistik timbulan sampah fasilitas.")
    add_screenshot_placeholder(doc, "SS-22", "Beranda PWA Tampilan Operator TPS3R", "Beranda Operator TPS3R.", "Smartphone Screen (9:16)")
    add_screenshot_placeholder(doc, "SS-23", "Formulir Input Sampah Masuk Fasilitas TPS3R", "Form input fasilitas TPS3R.", "Smartphone Screen (9:16)")

    add_heading_1(doc, "BAB 6 — PANDUAN UNTUK PERAN: PETUGAS ANGKUT / ARMADA")
    add_heading_2(doc, "6.1. Beranda Petugas Angkut & Monitoring Armada")
    add_body_p(doc, "Statistik ritase pengangkutan dan status kendaraan (Aktif / Pemeliharaan / Rusak).")
    add_screenshot_placeholder(doc, "SS-24", "Beranda PWA Tampilan Petugas Angkut", "Beranda Petugas Angkut.", "Smartphone Screen (9:16)")
    add_screenshot_placeholder(doc, "SS-25", "Halaman Monitoring Armada Kendaraan", "Halaman daftar armada.", "Mobile / Tablet Screen")

    add_heading_1(doc, "BAB 7 — PANDUAN UNTUK PERAN: KOORDINATOR LAPANGAN")
    add_heading_2(doc, "7.1. Beranda Koordinator & Validasi Data Lapangan")
    add_body_p(doc, "Wewenang kecamatan dan verifikasi berjenjang (Setujui / Tolak dengan catatan).")
    add_screenshot_placeholder(doc, "SS-26", "Beranda PWA Tampilan Koordinator Lapangan", "Beranda Koordinator.", "Smartphone Screen (9:16)")
    add_screenshot_placeholder(doc, "SS-27", "Halaman Validasi Data Lapangan (Daftar Antrean Pending)", "Halaman validasi data.", "Laptop / Tablet Screen (16:9)")
    add_screenshot_placeholder(doc, "SS-28", "Modal Konfirmasi Validasi Data (Setujui / Tolak dengan Catatan)", "Modal penolakan data.", "Mobile / Tablet Screen")
    add_screenshot_placeholder(doc, "SS-29", "Daftar Data Tervalidasi (Status Disetujui / Approved)", "Daftar data disetujui.", "Laptop Screen")

    add_heading_1(doc, "BAB 8 — PANDUAN UNTUK PERAN: EKSEKUTIF / KEPALA DINAS")
    add_heading_2(doc, "8.1. Dashboard Ringkasan Eksekutif")
    add_body_p(doc, "4 KPI Cards, Grafik Tren Volume Sampah Harian & Prediksi Machine Learning (Meta Prophet 7 hari ke depan dengan Disclaimer DSS), Grafik Komposisi SIPSN, dan Rekapitulasi Aduan.")
    add_screenshot_placeholder(doc, "SS-30", "Dashboard Eksekutif — 4 KPI Cards Metrik Utama", "4 KPI Cards Eksekutif.", "Desktop Full Width (16:9)")
    add_screenshot_placeholder(doc, "SS-31", "Grafik Tren Volume Sampah & Prediksi Machine Learning (Prophet)", "Grafik tren & prediksi ML Prophet.", "Desktop / Laptop Screen (16:9)")
    add_screenshot_placeholder(doc, "SS-32", "Grafik Komposisi Sampah Standar SIPSN", "Donut chart komposisi SIPSN.", "Desktop Screen")
    add_screenshot_placeholder(doc, "SS-33", "Ringkasan Rekapitulasi Aduan Masyarakat di Dashboard", "Widget rekapitulasi aduan.", "Desktop Screen")

    add_heading_2(doc, "8.2. Peta GIS (Distribusi Titik Sampah & Fasilitas)")
    add_body_p(doc, "Pemetaan spasial sebaran TPS, TPS3R, Bank Sampah, TPA Winong, dan titik aduan warga.")
    add_screenshot_placeholder(doc, "SS-34", "Peta GIS Distribusi Titik Sampah & Fasilitas Kebersihan", "Peta GIS interaktif.", "Desktop Full Width (16:9)")

    add_heading_1(doc, "BAB 9 — PANDUAN UNTUK PERAN: ADMINISTRATOR (SUPER ADMIN)")
    add_heading_2(doc, "9.1. Akses Navigasi & Manajemen Aduan Admin")
    add_body_p(doc, "Sidebar navigasi penuh, ubah status aduan, dan penulisan Tanggapan Dinas.")
    add_screenshot_placeholder(doc, "SS-35", "Sidebar Navigasi Dashboard Tampilan Super Admin", "Sidebar navigasi admin.", "Desktop Sidebar View")
    add_screenshot_placeholder(doc, "SS-36", "Halaman Manajemen Aduan Masyarakat (Daftar Kabupaten)", "Daftar aduan kabupaten.", "Laptop Screen (16:9)")
    add_screenshot_placeholder(doc, "SS-37", "Form Update Status Aduan & Penulisan Tanggapan Dinas", "Form update status & tanggapan dinas.", "Laptop Screen")

    add_heading_2(doc, "9.2. Modul Laporan & Master Data")
    add_body_p(doc, "Ekspor file Excel / SIPSN dan 6 tab Master Data (Wilayah, Lokasi, Armada, Kode Undangan, Penduduk, Pengguna).")
    add_screenshot_placeholder(doc, "SS-38", "Halaman Laporan & Modul Ekspor Data (Excel / SIPSN)", "Halaman ekspor laporan.", "Laptop Screen (16:9)")
    add_screenshot_placeholder(doc, "SS-39", "Master Data — Tab Manajemen Wilayah (Kecamatan & Desa)", "Tab master wilayah.", "Laptop Screen")
    add_screenshot_placeholder(doc, "SS-40", "Master Data — Tab Manajemen Lokasi & Fasilitas", "Tab master fasilitas.", "Laptop Screen")
    add_screenshot_placeholder(doc, "SS-41", "Master Data — Tab Manajemen Armada Kendaraan", "Tab master armada.", "Laptop Screen")
    add_screenshot_placeholder(doc, "SS-42", "Master Data — Tab Manajemen Kode Undangan Registrasi Petugas", "Tab master kode undangan.", "Laptop Screen")

    add_heading_2(doc, "9.3. Manajemen MoU, Intervensi Desa & AI SIMPAH Buddy")
    add_body_p(doc, "Manajemen MoU dengan indikator peringatan expired, skor intervensi desa, dan widget AI Chat Assistant Qwen LLM.")
    add_screenshot_placeholder(doc, "SS-43", "Halaman Manajemen MoU Kerjasama", "Daftar MoU kerja sama.", "Laptop Screen (16:9)")
    add_screenshot_placeholder(doc, "SS-44", "Formulir Tambah / Edit Dokumen MoU", "Form tambah MoU baru.", "Laptop Screen")
    add_screenshot_placeholder(doc, "SS-45", "Halaman Intervensi & Profil Pengelolaan Sampah Desa", "Halaman intervensi desa.", "Laptop Screen (16:9)")
    add_screenshot_placeholder(doc, "SS-46", "Halaman Manajemen Artikel Edukasi Publik", "Halaman edukasi admin.", "Laptop Screen")
    add_screenshot_placeholder(doc, "SS-47", "Tampilan Widget SIMPAH Buddy (AI Chat Assistant)", "Widget AI Chat SIMPAH Buddy.", "Mobile / Laptop Floating Widget View")

    # BAB 10
    add_heading_1(doc, "BAB 10 — PORTAL PUBLIK (TANPA LOGIN)")
    add_heading_2(doc, "10.1. Beranda Portal Publik")
    add_body_p(doc, "Portal Publik SIMPAH (www.simpah.id) dapat diakses bebas tanpa login, menyajikan landing page resmi, statistik publik, dan tombol Masuk Sistem.")
    add_screenshot_placeholder(doc, "SS-48", "Halaman Utama Portal Publik SIMPAH", "Landing page portal publik SIMPAH.", "Desktop Full Width (16:9)")

    add_heading_2(doc, "10.2. Tentang, Edukasi & Regulasi")
    add_body_p(doc, "Informasi pengembang STUDIO INKLUSI, galeri artikel edukasi pemilahan sampah, serta perundang-undangan dan regulasi kebersihan.")
    add_screenshot_placeholder(doc, "SS-49", "Halaman Galeri Artikel Edukasi Publik", "Halaman edukasi publik.", "Desktop / Laptop Screen")
    add_screenshot_placeholder(doc, "SS-50", "Halaman Regulasi & Peraturan Kebersihan", "Halaman regulasi kebersihan.", "Desktop Screen")

    add_heading_2(doc, "10.3. Pengaduan Publik & Pelacakan Resi")
    add_body_p(doc, "Pengiriman aduan langsung dari portal publik dan pelacakan status aduan via nomor resi dengan Progress Tracker 4 Tahap.")
    add_screenshot_placeholder(doc, "SS-51", "Form Pengaduan Masyarakat pada Portal Publik", "Form aduan publik tanpa login.", "Laptop / Tablet Screen")
    add_screenshot_placeholder(doc, "SS-52", "Halaman Cek Status & Pelacakan Resi Aduan Publik", "Pelacakan resi aduan publik.", "Laptop Screen (16:9)")

    # BAB 11
    add_heading_1(doc, "BAB 11 — FAQ & TROUBLESHOOTING")
    add_heading_2(doc, "11.1. Pertanyaan Umum (FAQ)")
    add_heading_3(doc, "Q1: Apakah SIMPAH perlu diunduh dari PlayStore?")
    add_body_p(doc, "A: Tidak perlu. SIMPAH adalah Progressive Web App (PWA) yang dapat diakses via browser dan diinstall langsung ke layar utama HP.")
    add_heading_3(doc, "Q2: Apakah pencatatan tetap bisa dilakukan saat offline?")
    add_body_p(doc, "A: Bisa. Data tersimpan di memori HP (IndexedDB) dan di-sync otomatis saat terhubung kembali ke internet.")

    add_heading_2(doc, "11.2. Masalah Umum & Solusi")
    add_bullet_p(doc, "Periksa jaringan internet, muat ulang halaman, atau gunakan tombol Reset System jika terjadi kendala cache.", "Troubleshooting: ")

    add_heading_2(doc, "11.3. Kontak Bantuan Teknis")
    add_body_p(doc, "Pengembang Sistem: STUDIO INKLUSI | Email: admin@simpah.id | Website: https://www.simpah.id")

    # LAMPIRAN
    add_heading_1(doc, "LAMPIRAN-LAMPIRAN")
    add_heading_2(doc, "Lampiran A: Matriks Hak Akses RBAC & RLS Security")
    add_body_p(doc, "Tabel pengamanan data berbasis Row Level Security (RLS) PostgreSQL Supabase.")

    add_heading_2(doc, "Lampiran B: Kategori Sampah Standar SIPSN")
    add_body_p(doc, "Definisi 8 kategori sampah standar pelaporan nasional KLHK.")

    add_heading_2(doc, "Lampiran C: Diagram Alur Status Pengaduan Masyarakat")
    add_body_p(doc, "Alur 4 tahap aduan warga (Baru -> Diproses -> Ditindaklanjuti -> Selesai / Ditolak).")
    add_screenshot_placeholder(doc, "SS-53", "Diagram Alur Perubahan Status Aduan Masyarakat", "Diagram alur status aduan warga.", "Diagram View (16:9)")

    add_heading_2(doc, "Lampiran D: Glosarium Istilah Teknis & Operasional")
    add_body_p(doc, "Glosarium istilah PWA, RLS, DSS, SIPSN, TPS3R, Ritase, Residu, LLM, Qwen.")

    # Save document safely
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully completed at primary path: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary file locked by Word. Document successfully saved at fallback path: {target_out}")

if __name__ == "__main__":
    create_document()
