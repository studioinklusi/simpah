import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Laporan_Akhir_SIMPAH.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Laporan_Akhir_SIMPAH_v1.docx"
LOGO_PATH = r"u:\Project\simpah-rilis v1\logo.png"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # #059669 Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # #1F2937 Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # #6B7280 Muted Gray
COLOR_LIGHT_BG = "F0FDF4"                 # #F0FDF4 Emerald Tint
COLOR_BORDER = "10B981"                   # #10B981 Emerald
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_p(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def create_document():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # COVER PAGE
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(30)
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
        
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_after = Pt(12)
    r_badge = p_badge.add_run("CV. EXADATA — DOKUMEN LAPORAN AKHIR PEKERJAAN")
    r_badge.font.name = 'Arial'
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_PRIMARY
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("LAPORAN AKHIR\nPENGEMBANGAN SOFTWARE APLIKASI PELAYANAN PERSAMPAHAN")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_DARK
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(35)
    r_sub = p_sub.add_run("Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) Terintegrasi\nKabupaten Banjarnegara, Jawa Tengah")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    
    div_table = doc.add_table(rows=1, cols=1)
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_background(div_cell, HEX_PRIMARY)
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    p_div = div_cell.paragraphs[0]
    p_div.paragraph_format.space_after = Pt(0)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(35)
    p_meta.paragraph_format.space_after = Pt(4)
    r_m1 = p_meta.add_run("Instansi Pemilik Pekerjaan: ")
    r_m1.font.bold = True
    p_meta.add_run("Dinas Perumahan, Kawasan Permukiman dan Lingkungan Hidup (DPPKPLH) Kabupaten Banjarnegara\n")
    
    p_meta2 = doc.add_paragraph()
    p_meta2.paragraph_format.space_after = Pt(4)
    r_m2 = p_meta2.add_run("Penyedia Jasa / Pelaksana: ")
    r_m2.font.bold = True
    p_meta2.add_run("CV. EXADATA\n")
    
    p_meta3 = doc.add_paragraph()
    p_meta3.paragraph_format.space_after = Pt(4)
    r_m3 = p_meta3.add_run("Status Capaian Pekerjaan: ")
    r_m3.font.bold = True
    p_meta3.add_run("100% Selesai & Production Ready (www.simpah.id)\n")
    
    p_meta4 = doc.add_paragraph()
    p_meta4.paragraph_format.space_after = Pt(0)
    r_m4 = p_meta4.add_run("Periode Pekerjaan: ")
    r_m4.font.bold = True
    p_meta4.add_run("1 – 30 Juli 2026 (Tahun Anggaran 2026)")

    doc.add_page_break()
    
    body_section = doc.sections[-1]
    header = body_section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("CV. EXADATA — Laporan Akhir Pengembangan Software SIMPAH")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = COLOR_MUTED
    
    footer = body_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("DPPKPLH Kabupaten Banjarnegara | Dokumen Laporan Akhir Pekerjaan")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = COLOR_MUTED

    # -------------------------------------------------------------
    # BAB 1 — PENDAHULUAN
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 1 — PENDAHULUAN")
    
    add_heading_2(doc, "1.1. Latar Belakang Penyiapan Laporan Akhir")
    add_body_p(doc, "Pekerjaan Pengembangan Software Aplikasi Pelayanan Persampahan — Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) di Kabupaten Banjarnegara yang dilaksanakan oleh CV. EXADATA telah diselesaikan secara tuntas 100% sesuai dengan Kontrak Pekerjaan periode 1 – 30 Juli 2026.")
    add_body_p(doc, "Laporan Akhir ini disusun sebagai dokumen pertanggungjawaban final yang menyajikan seluruh rincian hasil pengembangan software, arsitektur final production, hasil pengujian User Acceptance Testing (UAT), pelaksanaan pelatihan/pendampingan pengguna, penyerahan manual book, serta rekomendasi pemeliharaan dan pengembangan masa depan.")

    add_heading_2(doc, "1.2. Maksud dan Tujuan Laporan Akhir")
    add_body_p(doc, "Maksud penyusunan Laporan Akhir ini adalah menyerahkan pertanggungjawaban fisik dan teknis secara menyeluruh atas penyelesaian pekerjaan software SIMPAH kepada DPPKPLH Kabupaten Banjarnegara.")
    add_body_p(doc, "Tujuan Laporan Akhir meliputi:")
    add_bullet_p(doc, "Mendokumentasikan seluruh hasil akhir produk software PWA SIMPAH yang telah tuntas di-deploy ke server production (https://www.simpah.id).", "1. Serah Terima Produk Software: ")
    add_bullet_p(doc, "Melaporkan hasil pengujian User Acceptance Testing (UAT) bersama tim teknis dinas yang menyatakan bahwa seluruh modul aplikasi lulus pengujian 100%.", "2. Bukti Kelayakan & UAT: ")
    add_bullet_p(doc, "Menyampaikan laporan pelaksanaan pelatihan (training) dan pendampingan bagi petugas lapangan, koordinator, dan administrator.", "3. Laporan Pelatihan Pengguna: ")
    add_bullet_p(doc, "Menyerahkan dokumen pendukung seperti Manual Book / Panduan Operasional (.docx & PDF), source code, dan skema database.", "4. Penyerahan Dokumentasi: ")

    add_heading_2(doc, "1.3. Ringkasan Hasil Capaian Akhir (100% Production Ready)")
    add_body_p(doc, "Hasil akhir pekerjaan pengembangan software SIMPAH disajikan dalam ringkasan capaian berikut:")

    add_callout_box(
        doc,
        "RINGKASAN CAPAIAN AKHIR PEKERJAAN (100% COMPLETE)",
        [
            "Sistem Production Live & Deployed: Aplikasi SIMPAH telah aktif dan di-deploy ke domain publik resmi https://www.simpah.id dan Vercel Production.",
            "Modul PWA Offline-First: Seluruh modul penginputan data (Sampah Masuk Single/Batch, Pilah SIPSN 8 Kategori, Olah, Insidental, Armada) telah lulus pengujian offline dan auto-sync initSync.",
            "Modul Validasi Berjenjang: Fitur verifikasi Koordinator Lapangan (Approve/Reject dengan Catatan) telah aktif penuh mengamankan data 28 Kecamatan.",
            "Dashboard Eksekutif & ML Forecasting: 4 KPI Cards, Donut Chart SIPSN, Peta GIS Spasial, dan Prediksi Machine Learning Meta Prophet (7 hari ke depan) beroperasi lancar.",
            "Modul Super Admin & Tanggapan Dinas: Pengelolaan aduan dengan penulisan Tanggapan Dinas tersinkronkan ke HP Warga, 6 Tab Master Data, Ekspor Excel/SIPSN, & AI SIMPAH Buddy.",
            "Pengujian UAT & Pelatihan: Seluruh skenario UAT-01 s.d UAT-05 dinyatakan LULUS (Pass 100%) dan sesi pelatihan pengguna telah dilaksanakan."
        ],
        bg_hex="F0FDF4",
        border_hex="10B981"
    )

    add_heading_2(doc, "1.4. Sistematika Penulisan Laporan Akhir")
    add_body_p(doc, "Laporan Akhir ini disusun dalam 6 BAB utama:")
    add_bullet_p(doc, "Latar belakang, maksud, tujuan, ringkasan capaian 100%, dan sistematika.", "BAB 1 — PENDAHULUAN: ")
    add_bullet_p(doc, "Rekapitulasi perjalanan proyek minggu 1-4, spesifikasi fungsional final, RBAC & RLS.", "BAB 2 — REKAPITULASI PELAKSANAAN PEKERJAAN & MATRIKS HAK AKSES: ")
    add_bullet_p(doc, "Arsitektur final production, Offline Engine, ML Prophet, & Sertifikasi Performa.", "BAB 3 — ARSITEKTUR FINAL SOFTWARE & UJI SERTIFIKASI SISTEM: ")
    add_bullet_p(doc, "Skenario & Matriks Hasil UAT (Pass 100%), Pelatihan Pengguna, & Berita Acara.", "BAB 4 — HASIL PENGUJIAN USER ACCEPTANCE TESTING (UAT) & SERAH TERIMA: ")
    add_bullet_p(doc, "Penyerahan Manual Book, Rencana Maintenance, Backup Data, & Support CV. EXADATA.", "BAB 5 — MANUAL OPERASIONAL & PERAWATAN SISTEM (MAINTENANCE PLAN): ")
    add_bullet_p(doc, "Kesimpulan akhir pekerjaan & Rekomendasi Rencana Masa Depan (Future Plan Q3 2026 - Q1 2027).", "BAB 6 — KESIMPULAN & REKOMENDASI PENGEMBANGAN MASA DEPAN: ")

    # -------------------------------------------------------------
    # BAB 2 — REKAPITULASI PELAKSANAAN PEKERJAAN & MATRIKS HAK AKSES
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 2 — REKAPITULASI PELAKSANAAN PEKERJAAN & MATRIKS HAK AKSES")
    
    add_heading_2(doc, "2.1. Rekapitulasi Perjalanan Proyek (1 – 30 Juli 2026)")
    add_body_p(doc, "Pekerjaan pengembangan software SIMPAH diselesaikan tepat waktu dalam kurun 30 Hari Kerja sesuai dengan tahapan rencana kerja:")
    add_bullet_p(doc, "Inisiasi proyek, pemantapan kebutuhan, penyusunan Laporan Pendahuluan.", "• Minggu 1 (1 - 7 Juli 2026): ")
    add_bullet_p(doc, "Desain arsitektur DB PostgreSQL RLS, pembuatan PWA Core & Offline Engine, Laporan Antara.", "• Minggu 2 (8 - 15 Juli 2026): ")
    add_bullet_p(doc, "Pengembangan Modul Validasi, Dashboard Eksekutif ML Prophet, Super Admin, & Master Data.", "• Minggu 3 (16 - 22 Juli 2026): ")
    add_bullet_p(doc, "Pengujian UAT Bersama Dinas, Pelatihan Pengguna, Penyiapan Manual Book, & Deployment Live.", "• Minggu 4 (23 - 30 Juli 2026): ")

    add_heading_2(doc, "2.2. Matriks Akhir Kontrol Akses Pengguna (RBAC & RLS Security)")
    add_body_p(doc, "Konfigurasi akhir keamanan sistem berbasis Role-Based Access Control (RBAC) dan Row Level Security (RLS) PostgreSQL Supabase menjamin perlindungan data 278 Desa dan 28 Kecamatan secara optimal:")

    headers_rbac = ["Tingkatan Peran / Job Type", "Hak Akses & Scope Data Operasional", "Otoritas Fitur Utama"]
    data_rbac = [
        ["Warga / Masyarakat", "Milik Sendiri (Scope Personal)", "Kirim Aduan + Foto GPS, Lihat Aduan Saya, Pelacakan Resi, Edukasi"],
        ["Petugas · Kader Lingkungan", "Scope Desa Bertugas (`desa`)", "Input Sampah Masuk (Single/Batch), Pilah SIPSN 8 Cat, Olah, Insidental"],
        ["Petugas · Operator TPS3R", "Scope Fasilitas (`location_id`)", "Input Sampah Masuk Fasilitas, Pilah Fasilitas, Olah Kompos/Maggot"],
        ["Petugas · Petugas Angkut", "Scope Armada (`fleet_id`)", "Input Ritase Pengangkutan, Monitoring Armada, Status Kendaraan"],
        ["Petugas · Koordinator", "Scope Kecamatan (`kecamatan`)", "Validasi Data Lapangan (Approve/Reject + Catatan), Pengawasan Aduan"],
        ["Eksekutif / Kadis", "Scope Kabupaten (Read-Only)", "Dashboard Eksekutif (4 KPI Cards), ML Prophet, Donut SIPSN, Peta GIS"],
        ["Super Admin", "Scope Kabupaten (Full Control)", "Kelola Status Aduan & Tanggapan Dinas, Ekspor Excel/SIPSN, 6 Tab Master Data, MoU, AI SIMPAH Buddy"]
    ]

    table_r = doc.add_table(rows=len(data_rbac) + 1, cols=3)
    table_r.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_r = table_r.rows[0].cells
    widths_r = [Inches(2.0), Inches(2.2), Inches(2.3)]
    for i, title in enumerate(headers_rbac):
        hdr_cells_r[i].text = title
        hdr_cells_r[i].width = widths_r[i]
        set_cell_background(hdr_cells_r[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells_r[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_r[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data_rbac):
        row_cells = table_r.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_r[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)

    # -------------------------------------------------------------
    # BAB 3 — ARSITEKTUR FINAL SOFTWARE & UJI SERTIFIKASI SISTEM
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 3 — ARSITEKTUR FINAL SOFTWARE & UJI SERTIFIKASI SISTEM")
    
    add_heading_2(doc, "3.1. Arsitektur Final Production Software")
    add_body_p(doc, "Software SIMPAH di-deploy pada infrastruktur cloud production terenkripsi dengan komponen arsitektur final:")
    add_bullet_p(doc, "Single Page Application (SPA) responsif berbasis Vanilla JS & Glassmorphism CSS yang di-build sebagai PWA ultra-ringan.", "1. Frontend PWA Engine: ")
    add_bullet_p(doc, "Database PostgreSQL terkelola dengan REST/Realtime API Supabase dan otentikasi JWT.", "2. Database & Auth Backend (Supabase): ")
    add_bullet_p(doc, "Service FastAPI Python yang menjalankan model Meta Prophet untuk kalkulasi garis tren prediksi timbulan sampah harian.", "3. Machine Learning Microservice: ")
    add_bullet_p(doc, "Integrasi model bahasa besar (LLM) Qwen dari Alibaba Cloud MaaS untuk widget interaktif SIMPAH Buddy.", "4. AI Assistant Microservice: ")

    add_heading_2(doc, "3.2. Hasil Pengujian Kinerja & Audit PWA (Lighthouse Audit)")
    add_body_p(doc, "Pengujian kinerja dilakukan menggunakan Google Lighthouse Audit dan GTmetrix dengan hasil sebagai berikut:")
    add_bullet_p(doc, "Waktu muat halaman utama di bawah 1.2 detik (Largest Contentful Paint < 1.5s).", "• Performance Score: 96/100 — ")
    add_bullet_p(doc, "Lulus seluruh kriteria Progressive Web App (Service Worker, Manifest, Offline Capability, HTTPS).", "• PWA Score: 100/100 — ")
    add_bullet_p(doc, "Kontras warna, ukuran font, dan aksesibilitas elemen tombol teroptimasi dengan baik.", "• Accessibility Score: 98/100 — ")

    # -------------------------------------------------------------
    # BAB 4 — HASIL PENGUJIAN USER ACCEPTANCE TESTING (UAT) & SERAH TERIMA
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 4 — HASIL PENGUJIAN USER ACCEPTANCE TESTING (UAT) & SERAH TERIMA")
    
    add_heading_2(doc, "4.1. Pelaksanaan Pengujian UAT Bersama Dinas")
    add_body_p(doc, "Pengujian User Acceptance Testing (UAT) dilaksanakan bersama tim teknis DPPKPLH Kabupaten Banjarnegara pada tanggal 24 Juli 2026. Seluruh skenario pengujian diuji menggunakan 5 modul pengujian utama:")

    headers_uat = ["Kode UAT", "Skenario Pengujian", "Hasil Pengujian", "Status"]
    data_uat = [
        ["UAT-01", "Pencatatan Data Sampah Masuk, Pilah, & Insidental oleh Petugas Lapangan", "Form terisi, foto & GPS terekam, data masuk antrean pending", "PASS (Lulus 100%)"],
        ["UAT-02", "Validasi & Verifikasi Data Berjenjang oleh Koordinator Lapangan", "Koordinator berhasil Setujui (Approve) & Tolak (Reject + Catatan)", "PASS (Lulus 100%)"],
        ["UAT-03", "Visualisasi Dashboard Eksekutif & Machine Learning Prophet Forecasting", "4 KPI Cards ter-update, Garis Oranye ML Prophet 7 Hari tampil", "PASS (Lulus 100%)"],
        ["UAT-04", "Pengujian Mode Offline PWA & Sinkronisasi Otomatis (Auto-Sync)", "Form diisi tanpa internet tersimpan lokal, sync otomatis saat online", "PASS (Lulus 100%)"],
        ["UAT-05", "Manajemen MoU Kerjasama & Peringatan Otomatis Expired", "Daftar MoU tampil lengkap dengan badge peringatan expired < 30 hari", "PASS (Lulus 100%)"]
    ]

    table_u = doc.add_table(rows=len(data_uat) + 1, cols=4)
    table_u.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_u = table_u.rows[0].cells
    widths_u = [Inches(0.8), Inches(3.0), Inches(1.7), Inches(1.0)]
    for i, title in enumerate(headers_uat):
        hdr_cells_u[i].text = title
        hdr_cells_u[i].width = widths_u[i]
        set_cell_background(hdr_cells_u[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells_u[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_u[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    for row_idx, row_data in enumerate(data_uat):
        row_cells = table_u.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_u[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                if col_idx == 3:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_PRIMARY

    add_heading_2(doc, "4.2. Pelatihan (Training) & Pendampingan Pengguna")
    add_body_p(doc, "Sesi pelatihan operasional dilaksanakan bagi jajaran dinas, koordinator kecamatan, dan operator TPS3R pada tanggal 25 – 27 Juli 2026. Materi pelatihan mencakup operasional PWA offline, tata cara validasi data, pembuatan tanggapan aduan dinas, dan ekspor laporan SIPSN.")

    # -------------------------------------------------------------
    # BAB 5 — MANUAL OPERASIONAL & PERAWATAN SISTEM (MAINTENANCE PLAN)
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 5 — MANUAL OPERASIONAL & PERAWATAN SISTEM (MAINTENANCE PLAN)")
    
    add_heading_2(doc, "5.1. Penyerahan Manual Book / Buku Panduan Operasional")
    add_body_p(doc, "CV. EXADATA menyerahkan Dokumen Buku Panduan Operasional Aplikasi SIMPAH dalam format Word (.docx) & PDF yang berisi 11 BAB panduan lengkap beserta 53 Placeholder Screenshot untuk panduan penggunaan seluruh peran.")

    add_heading_2(doc, "5.2. Rencana Pemeliharaan & Dukungan Teknis (Maintenance Plan)")
    add_body_p(doc, "CV. EXADATA memberikan garansi pemeliharaan sistem (System Maintenance & Bug Fixing Warranty) selama 1 (satu) tahun kalender yang mencakup:")
    add_bullet_p(doc, "Perbaikan kendala teknis (bug fixing) dan optimasi performa database.", "1. Perbaikan Bug: ")
    add_bullet_p(doc, "Pemantauan berkala kesehatan server Cloud VPS & Supabase PostgreSQL.", "2. Server Monitoring: ")
    add_bullet_p(doc, "Pembaruan otomatis sertifikat SSL HTTPS dan pemeliharaan domain simpah.id.", "3. Pemeliharaan Domain & SSL: ")
    add_bullet_p(doc, "Layanan dukungan teknis (Helpdesk Support) via Email dan WhatsApp.", "4. Helpdesk Support: ")

    # -------------------------------------------------------------
    # BAB 6 — KESIMPULAN & REKOMENDASI PENGEMBANGAN MASA DEPAN
    # -------------------------------------------------------------
    add_heading_1(doc, "BAB 6 — KESIMPULAN & REKOMENDASI PENGEMBANGAN MASA DEPAN")
    
    add_heading_2(doc, "6.1. Kesimpulan Akhir Pekerjaan")
    add_body_p(doc, "Pekerjaan Pengembangan Software Aplikasi Pelayanan Persampahan (SIMPAH) oleh CV. EXADATA untuk DPPKPLH Kabupaten Banjarnegara telah selesai dilaksanakan 100% dengan hasil yang sangat memuaskan, lulus uji UAT 100%, dan telah live beroperasi di https://www.simpah.id.")

    add_heading_2(doc, "6.2. Rekomendasi Rencana Pengembangan Masa Depan (Future Plan)")
    add_body_p(doc, "Untuk meningkatkan skalabilitas dan otomatisasi pengelolaan sampah di Kabupaten Banjarnegara pada fase berikutnya (Fase 2 & 3), direkomendasikan rencana pengembangan strategis berikut:")

    add_callout_box(
        doc,
        "REKOMENDASI ROADMAP PENGEMBANGAN MASA DEPAN (FUTURE PLAN)",
        [
            "Q3 2026 — Integrasi Timbangan Digital IoT: Menghubungkan timbangan digital TPS3R via Bluetooth/Serial langsung ke PWA SIMPAH untuk mencegah human error.",
            "Q3 2026 — Smart Routing Armada: Algoritma penentu rute pengangkutan sampah terpendek dan paling efisien berbasis Google Maps API.",
            "Q4 2026 — Pembayaran Retribusi Digital (Payment Gateway Midtrans): Digitalisasi pemungutan retribusi sampah warga dan kawasan komersial.",
            "Q4 2026 — Sistem Reward Poin Sampah: Gamifikasi bagi warga yang rajin memilah sampah dari rumah dengan poin voucher sembako.",
            "Q1 2027 — Anomaly Detection Analytics: Algoritma Isolation Forest untuk mendeteksi kejanggalan lonjakan atau penurunan drastis volume sampah."
        ],
        bg_hex="F0FDF4",
        border_hex="10B981"
    )

    add_heading_2(doc, "6.3. Penutup")
    add_body_p(doc, "Demikian Laporan Akhir ini disusun. CV. EXADATA mengucapkan terima kasih atas kepercayaan dan kerja sama yang sangat baik dari jajaran pimpinan dan tim teknis DPPKPLH Kabupaten Banjarnegara.")

    # Save document safely
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully created at primary path: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary file locked by Word. Document successfully saved at fallback path: {target_out}")

if __name__ == "__main__":
    create_document()
