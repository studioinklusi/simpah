import os
import sys

# Ensure python-docx is installed
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Installing python-docx...")
    os.system(f'"{sys.executable}" -m pip install python-docx')
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def add_heading(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(31, 78, 121)  # Dark Blue
        if level == 1:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(12)
            run.font.bold = True
    return heading

def add_paragraph(doc, text, bold=False, italic=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def main():
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Document Header
    p_header = add_paragraph(doc, "RENCANA ANGGARAN BIAYA (RAB)", bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=2)
    p_header.runs[0].font.size = Pt(16)
    p_header.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    
    p_sub = add_paragraph(doc, "SISTEM INFORMASI MANAJEMEN PENGELOLAAN SAMPAH (SIMPAH)", bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=18)
    p_sub.runs[0].font.size = Pt(12)
    p_sub.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    
    # Metadata Table / Section
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.autofit = False
    table_meta.columns[0].width = Inches(1.8)
    table_meta.columns[1].width = Inches(4.7)
    
    metadata = [
        ("Nama Inovasi", ": SIMPAH (Sistem Informasi Manajemen Pengelolaan Sampah)"),
        ("Instansi", ": Dinas Perumahan, Kawasan Permukiman dan Lingkungan Hidup (DPPKPLH)"),
        ("Periode", ": 1 Tahun (Anggaran Operasional & Pemeliharaan)"),
        ("Total Anggaran", ": Rp 50.000.000,- (Lima Puluh Juta Rupiah) - Termasuk PPN 11%")
    ]
    
    for idx, (label, val) in enumerate(metadata):
        # label
        cell_l = table_meta.rows[idx].cells[0]
        cell_l.text = ""
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(2)
        run_l = p_l.add_run(label)
        run_l.font.name = 'Arial'
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        
        # value
        cell_v = table_meta.rows[idx].cells[1]
        cell_v.text = ""
        p_v = cell_v.paragraphs[0]
        p_v.paragraph_format.space_after = Pt(2)
        run_v = p_v.add_run(val)
        run_v.font.name = 'Arial'
        run_v.font.size = Pt(10.5)
        if "Total" in label:
            run_v.font.bold = True
            run_v.font.color.rgb = RGBColor(180, 50, 50)
            
    doc.add_paragraph("\n")
    
    # Table Headings
    headers = ["No", "Komponen Pekerjaan / Rincian", "Keterangan / Spesifikasi", "Vol", "Satuan", "Harga Satuan", "Total Harga"]
    
    # Create main RAB table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Widths of columns in inches
    col_widths = [Inches(0.4), Inches(2.2), Inches(2.0), Inches(0.4), Inches(0.6), Inches(1.1), Inches(1.1)]
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(title_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "1F4E79") # Dark Blue
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        hdr_cells[i].width = col_widths[i]
        
    items_data = [
        # 1. LISENSI
        ("1", "Lisensi Penggunaan Sistem (1 Tahun)", "", "", "", "", "22.045.000", True),
        ("1.1", "Lisensi Platform Core SIMPAH (PWA)", "Akses source code, modul aplikasi web responsif (Warga, Petugas, Koordinator, Eksekutif) selama 1 tahun", "1", "Paket", "22.045.000", "22.045.000", False),
        ("Subtotal 1", "Subtotal Lisensi Penggunaan Sistem", "", "", "", "", "22.045.000", True, "F2F2F2"),
        
        # 2. KUSTOMISASI
        ("2", "Kustomisasi & Setup Database PostgreSQL", "", "", "", "", "12.000.000", True),
        ("2.1", "Pembuatan & Kustomisasi Skema Database", "Setup database PostgreSQL lokal di VPS (konfigurasi tabel, view rekapitulasi, & triggers)", "1", "Paket", "6.000.000", "6.000.000", False),
        ("2.2", "Penyesuaian API & Dashboard Dinas", "Kustomisasi form pelaporan sampah sesuai standar Dinas DPPKPLH & integrasi Peta GIS", "1", "Paket", "6.000.000", "6.000.000", False),
        ("Subtotal 2", "Subtotal Kustomisasi & Setup Database", "", "", "", "", "12.000.000", True, "F2F2F2"),
        
        # 3. SERVER & CLOUD
        ("3", "Sewa Server Cloud Database & Hosting (1 Tahun)", "", "", "", "", "6.000.000", True),
        ("3.1", "Sewa Server Cloud VPS (Virtual Private Server)", "Penyewaan Cloud VPS Server (Spesifikasi Core: 4 vCPU, 8GB RAM, 100GB SSD, Bandwidth Unlimited, Backup)", "12", "Bulan", "450.000", "5.400.000", False),
        ("3.2", "Domain Resmi Pemda & SSL Certificate", "Pengurusan domain .go.id / .id dan sertifikat SSL keamanan HTTPS untuk 1 tahun", "1", "Paket", "600.000", "600.000", False),
        ("Subtotal 3", "Subtotal Server Cloud & Hosting", "", "", "", "", "6.000.000", True, "F2F2F2"),
        
        # 4. INSTALASI & MANUAL
        ("4", "Instalasi, Uji Coba, Pemeliharaan & Manual Book", "", "", "", "", "5.000.000", True),
        ("4.1", "Setup CI/CD & Deployment Server", "Konfigurasi pipeline deployment, pengujian keamanan data, dan konfigurasi environment", "1", "Paket", "2.000.000", "2.000.000", False),
        ("4.2", "DevOps & Pemeliharaan Sistem (1 Tahun)", "Monitoring keaktifan server 24/7, perbaikan bug minor, dan optimasi database berkala", "1", "Paket", "1.500.000", "1.500.000", False),
        ("4.3", "Penyusunan Buku Panduan (Manual Book)", "Penyusunan modul manual cetak dan video tutorial penggunaan sistem untuk admin & kader", "1", "Paket", "1.500.000", "1.500.000", False),
        ("Subtotal 4", "Subtotal Instalasi, Pemeliharaan & Manual", "", "", "", "", "5.000.000", True, "F2F2F2"),
    ]
    
    for row_data in items_data:
        no, comp, spec, vol, sat, price, total, is_bold = row_data[:8]
        bg_color = row_data[8] if len(row_data) > 8 else None
        
        row_cells = table.add_row().cells
        
        # Populate cells
        data_row = [no, comp, spec, vol, sat, price, total]
        alignments = [
            WD_PARAGRAPH_ALIGNMENT.CENTER, # No
            WD_PARAGRAPH_ALIGNMENT.LEFT,   # Komponen
            WD_PARAGRAPH_ALIGNMENT.LEFT,   # Spesifikasi
            WD_PARAGRAPH_ALIGNMENT.CENTER, # Vol
            WD_PARAGRAPH_ALIGNMENT.CENTER, # Satuan
            WD_PARAGRAPH_ALIGNMENT.RIGHT,  # Harga
            WD_PARAGRAPH_ALIGNMENT.RIGHT   # Total
        ]
        
        for idx, text in enumerate(data_row):
            row_cells[idx].text = ""
            p = row_cells[idx].paragraphs[0]
            p.alignment = alignments[idx]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = is_bold
            
            if is_bold and no.startswith("Subtotal"):
                run.font.color.rgb = RGBColor(31, 78, 121)
                
            set_cell_margins(row_cells[idx], top=60, bottom=60, left=80, right=80)
            if bg_color:
                set_cell_background(row_cells[idx], bg_color)
            elif is_bold and not no.startswith("Subtotal"):
                set_cell_background(row_cells[idx], "E9EDF4") # Accent background for category title
                
            row_cells[idx].width = col_widths[idx]
            
    # Add Subtotal Akhir, PPN 11%, and Total
    totals_data = [
        ("SUB-TOTAL KESELURUHAN (Sebelum PPN)", "45.045.000", "F2F2F2"),
        ("PAJAK PERTAMBAHAN NILAI (PPN 11%)", "4.955.000", "F2F2F2"),
        ("TOTAL ESTIMASI BIAYA (Termasuk PPN)", "50.000.000", "D9E1F2")
    ]
    
    for title, val, bg_color in totals_data:
        row_cells = table.add_row().cells
        
        # Merge first 5 cells
        row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
        
        # Title Cell
        row_cells[0].text = ""
        p_title = row_cells[0].paragraphs[0]
        p_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p_title.paragraph_format.space_before = Pt(4)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run(title)
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(9.5)
        run_title.font.bold = True
        if "TOTAL" in title:
            run_title.font.color.rgb = RGBColor(31, 78, 121)
            
        # Empty cells width adjustments
        row_cells[5].text = ""
        
        # Value Cell
        row_cells[6].text = ""
        p_val = row_cells[6].paragraphs[0]
        p_val.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p_val.paragraph_format.space_before = Pt(4)
        p_val.paragraph_format.space_after = Pt(4)
        run_val = p_val.add_run(val)
        run_val.font.name = 'Arial'
        run_val.font.size = Pt(9.5)
        run_val.font.bold = True
        if "TOTAL" in title:
            run_val.font.color.rgb = RGBColor(180, 50, 50)
            
        # Apply padding and backgrounds
        for idx in range(7):
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=80, right=80)
            set_cell_background(row_cells[idx], bg_color)

    # Apply widths to all rows to prevent word resizing
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Text footer
    doc.add_paragraph("\n")
    p_terbilang = add_paragraph(doc, "Terbilang: Lima Puluh Juta Rupiah", bold=True, italic=True)
    p_terbilang.paragraph_format.space_before = Pt(6)
    
    add_paragraph(doc, "\nCatatan Teknis:")
    notes = [
        "Server utama dideploy menggunakan Cloud VPS Server (self-hosted) dengan spesifikasi yang andal untuk meng-host frontend PWA, backend web server Nginx, dan database PostgreSQL secara mandiri.",
        "Sistem dilengkapi dengan skema backup harian otomatis lokal/external dan SSL HTTPS demi menjamin kedaulatan data dan keamanan operasional.",
        "Pelatihan operasional dan sosialisasi kepada kader lingkungan serta petugas lapangan diselenggarakan dan ditanggung sepenuhnya oleh instansi pengguna (Dinas). Pihak pengembang menyediakan manual book (buku panduan) dan video tutorial penggunaan sistem sebagai bagian dari serah terima sistem."
    ]
    for note in notes:
        p_n = doc.add_paragraph(style='List Bullet')
        p_n.paragraph_format.space_after = Pt(3)
        run_n = p_n.add_run(note)
        run_n.font.name = 'Arial'
        run_n.font.size = Pt(9.5)

    # Signatures
    doc.add_paragraph("\n\n")
    table_sig = doc.add_table(rows=5, cols=2)
    table_sig.autofit = True
    
    # Left signature
    cell_l = table_sig.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    run_l = p_l.add_run("Mengetahui,\nKepala Dinas Perumahan, Kawasan Permukiman\ndan Lingkungan Hidup Kabupaten Banjarnegara\n\n\n\n\n(.......................................................)")
    run_l.font.name = 'Arial'
    run_l.font.size = Pt(10)
    
    # Right signature
    cell_r = table_sig.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_r = p_r.add_run("Dibuat Oleh,\nTim Pengembang Aplikasi SIMPAH\nBanjarnegara\n\n\n\n\n(.......................................................)")
    run_r.font.name = 'Arial'
    run_r.font.size = Pt(10)
    
    # Remove borders from signature table
    for row in table_sig.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'none')
                tc_borders.append(node)
            tc_pr.append(tc_borders)

    # Save path
    docx_path = r"u:\Project\simpah-rilis v1\docs\RAB_SIMPAH_50Juta.docx"
    doc.save(docx_path)
    print(f"RAB Word document saved successfully to: {docx_path}")
    
    # Try converting to PDF using win32com if on Windows with Word installed
    pdf_path = r"u:\Project\simpah-rilis v1\docs\RAB_SIMPAH_50Juta.pdf"
    try:
        import win32com.client
        print("Attempting to convert DOCX to PDF via MS Word...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_pdf = word.Documents.Open(os.path.abspath(docx_path))
        doc_pdf.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 is wdFormatPDF
        doc_pdf.Close()
        word.Quit()
        print(f"RAB PDF document generated successfully at: {pdf_path}")
    except Exception as e:
        print(f"Note: PDF generation skipped (either MS Word or win32com is not available). Error details: {e}")
        print("The DOCX and MD versions of the RAB are fully complete and ready.")

if __name__ == "__main__":
    main()
