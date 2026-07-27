import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Lampiran_C_Diagram_Alur_Status_Pengaduan_Masyarakat.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Lampiran_C_Diagram_Alur_Status_Pengaduan_Masyarakat_v1.docx"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # Muted Gray
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"
HEX_HEADER_TRANSITION = "1E293B"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📢 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def build_document():
    doc = docx.Document()
    
    # Page setup - Margins 0.75 inch
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # HEADER DOCUMENT
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("LAMPIRAN C: DIAGRAM ALUR STATUS PENGADUAN MASYARAKAT")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(17)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_s = p_sub.add_run("Sistem Informasi Manajemen Pengelolaan Sampah (SIMPAH) Kabupaten Banjarnegara")
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(11.5)
    r_s.font.italic = True
    r_s.font.color.rgb = COLOR_MUTED

    # OVERVIEW CALLOUT
    add_callout_box(doc, "Prinsip Kerja & State Machine Pengaduan Warga", [
        "Transparansi Pelaporan: Warga mendapatkan Nomor Resi Unik (misal: ADU-20260727-9A4B) untuk melacak status penanganan aduan secara real-time tanpa perlu login.",
        "Aturan Transisi Searah (State Machine): Status pengaduan bergerak searah (Baru -> Diproses -> Ditindaklanjuti -> Selesai / Ditolak) tanpa ada rollback ilegal.",
        "Akuntabilitas & SLA: Setiap perubahan status mencatat timestamp, penanggung jawab, serta foto bukti pengerjaan sebelum dan sesudah."
    ])

    # BAGIAN 1: TABEL TAHAPAN ADUAN
    add_heading_1(doc, "Bagian 1: Detail 4 Tahap Utama Alur Penanganan Pengaduan")
    add_body_p(doc, "Rincian alur kerja operasional dari pelaporan awal masyarakat hingga penutupan aduan oleh dinas:")

    headers_flow = ["Tahap Status", "Badge Warna", "Penanggung Jawab", "Kriteria Pemicu (Trigger)", "Aksi System & Input Data", "Target SLA", "Output Bagi Warga"]
    data_flow = [
        ["1. Baru", "Blue #3b82f6", "Warga / Pengunjung", "Warga mengisi form aduan via PWA/Portal (Foto + GPS).", "Generating Resi Unik, Notifikasi ke Koordinator Lapangan.", "< 1 Jam", "Nomor Resi Pelacakan Publik."],
        ["2. Diproses", "Amber #f59e0b", "Koordinator / Admin", "Koordinator memvalidasi keaslian lokasi & foto aduan.", "Ubah status ke Diproses, alokasi lokasi aduan ke Armada/Kecamatan.", "1 - 3 Jam", "Status resi: Diproses & Verifikasi Tim."],
        ["3. Ditindaklanjuti", "Purple #8b5cf6", "Petugas / Driver Armada", "Tim armada berangkat ke lokasi aduan & melakukan pembersihan.", "Ubah status ke Ditindaklanjuti, upload Foto Bukti Pengerjaan.", "3 - 12 Jam", "Warga melihat foto progress di lokasi."],
        ["4. Selesai", "Green #10b981", "Koordinator / Admin", "Koordinator mengonfirmasi lokasi sudah bersih 100%.", "Ubah status ke Selesai (Terminal), input Tanggapan Resmi Dinas.", "< 24 Jam", "Resi ditutup, foto Sebelum/Sesudah & Catatan Dinas."],
        ["Ditolak", "Red #ef4444", "Koordinator / Admin", "Laporan terdeteksi palsu, duplikat, atau di luar wewenang Pemkab.", "Ubah status ke Ditolak (Terminal), wajib input Alasan Penolakan.", "< 6 Jam", "Resi ditutup, transparansi alasan penolakan."]
    ]

    table_flow = doc.add_table(rows=len(data_flow) + 1, cols=7)
    table_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_flow = [Inches(1.0), Inches(0.9), Inches(1.0), Inches(1.1), Inches(1.2), Inches(0.6), Inches(1.0)]

    hdr_cells = table_flow.rows[0].cells
    for i, title in enumerate(headers_flow):
        hdr_cells[i].text = title
        hdr_cells[i].width = widths_flow[i]
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=90, bottom=90, left=60, right=60)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 1, 5] else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data_flow):
        row_cells = table_flow.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_flow[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=60, right=60)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 1, 5] else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                if col_idx == 0:
                    run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # BAGIAN 2: STATE MACHINE TRANSITION MATRIX
    add_heading_1(doc, "Bagian 2: Matriks Aturan Transisi Status Engine (VALID_TRANSITIONS)")
    add_body_p(doc, "Tabel aturan transisi state machine pada file src/pages/dashboard/aduan.js untuk mencegah lompatan status yang tidak valid:")

    headers_tr = ["Status Asal (Current State)", "Opsi Status Tujuan yang Diizinkan", "Sifat State & Pembatasan Rule"]
    data_tr = [
        ["baru", "diproses, ditolak", "State Awal. Wajib diverifikasi sebelum disposisi ke lapangan."],
        ["diproses", "ditindaklanjuti, ditolak", "State Verifikasi. Lokasi aduan dialokasikan ke armada/kader."],
        ["ditindaklanjuti", "selesai, ditolak", "State Pengerjaan. Wajib unggah foto bukti pembersihan."],
        ["selesai", "(Tidak ada - Terminal State)", "State Terkunci. Aduan berhasil ditutup dengan Tanggapan Dinas."],
        ["ditolak", "(Tidak ada - Terminal State)", "State Terkunci. Aduan ditolak resmi dengan Alasan Penolakan."]
    ]

    tbl_tr = doc.add_table(rows=len(data_tr) + 1, cols=3)
    tbl_tr.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_tr = [Inches(1.5), Inches(2.0), Inches(3.3)]

    hdr_tr = tbl_tr.rows[0].cells
    for i, title in enumerate(headers_tr):
        hdr_tr[i].text = title
        hdr_tr[i].width = widths_tr[i]
        set_cell_background(hdr_tr[i], HEX_HEADER_TRANSITION)
        set_cell_margins(hdr_tr[i], top=90, bottom=90, left=100, right=100)
        p = hdr_tr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data_tr):
        row_cells = tbl_tr.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_tr[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                if col_idx == 0:
                    run.font.bold = True

    # SAVE DOCUMENT
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully saved at: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary locked. Document saved at fallback: {target_out}")

if __name__ == "__main__":
    build_document()
