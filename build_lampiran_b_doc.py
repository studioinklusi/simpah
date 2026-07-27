import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOC_PATH = r"u:\Project\simpah-rilis v1\Lampiran_B_Kategori_Sampah_Standar_SIPSN.docx"
DOC_FALLBACK_PATH = r"u:\Project\simpah-rilis v1\Lampiran_B_Kategori_Sampah_Standar_SIPSN_v1.docx"

COLOR_PRIMARY = RGBColor(5, 150, 105)     # Emerald Green
COLOR_DARK = RGBColor(31, 41, 55)        # Dark Slate
COLOR_MUTED = RGBColor(107, 114, 128)    # Muted Gray
HEX_PRIMARY = "059669"
HEX_LIGHT_ROW = "F9FAFB"
HEX_HEADER_SPECIAL = "1E293B"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_body_p(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK
    run.bold = bold
    run.italic = italic
    return p

def add_callout_box(doc, title, content_list, bg_hex="F0FDF4", border_hex="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'color': border_hex, 'sz': '24'},
                    top={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    right={'val': 'single', 'color': 'D1D5DB', 'sz': '4'},
                    bottom={'val': 'single', 'color': 'D1D5DB', 'sz': '4'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"🍃 {title}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_PRIMARY
    
    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        r = p_item.add_run(f"• {item}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_DARK
        
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def build_document():
    doc = docx.Document()
    
    # Page setup - Margins 0.75 inch
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # HEADER DOCUMENT
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("LAMPIRAN B: KATEGORI SAMPAH STANDAR SIPSN")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(18)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_s = p_sub.add_run("Standar Pelaporan Nasional Kementerian Lingkungan Hidup dan Kehutanan (KLHK)")
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(12)
    r_s.font.italic = True
    r_s.font.color.rgb = COLOR_MUTED

    # OVERVIEW CALLOUT
    add_callout_box(doc, "Standar Klasifikasi Pelaporan SIPSN KLHK", [
        "Sistem Informasi Pengelolaan Sampah Nasional (SIPSN) KLHK mewajibkan pengelompokan komposisi sampah secara terstandar di seluruh Indonesia.",
        "Aplikasi SIMPAH mengadopsi 8 Kategori Utama SIPSN untuk perhitungan Recycling Rate, komposisi timbulan harian, dan agregasi laporan otomatis ke format CSV/Excel SIPSN.",
        "Metode agregasi menghitung tonase harian: Total Ton = (Total Weight kg) / 1000."
    ])

    # BAGIAN 1: TABLE KATEGORI SAMPAH
    add_heading_1(doc, "Bagian 1: Matriks Definisi 8 Kategori Sampah Standar SIPSN KLHK")
    add_body_p(doc, "Definisi baku 8 kategori sampah standar nasional beserta contoh jenis barang dan potensi pengolahannya di SIMPAH Engine:")

    headers_sipsn = ["No", "Kode", "Kategori Sampah SIPSN", "Definisi & Karakteristik Material", "Contoh Jenis Barang / Limbah", "Potensi Metode Pengolahan (SIMPAH)"]
    data_sipsn = [
        ["1", "SM", "Sisa Makanan", "Sampah organik basah sisa olahan makanan & konsumsi kuliner.", "Nasi basi, sisa sayuran, buah busuk, sisa daging/ikan, ampas kopi.", "Biokonversi Maggot BSF, Pengomposan, Pakan Ternak, Biogas, Eco-Enzyme."],
        ["2", "KR", "Kayu / Ranting", "Sampah organik kering bertestur keras pemangkasan kebun & kayu.", "Dedaunan kering, ranting pohon, sisa potongan kayu, bambu, serbuk gergaji.", "Pencacah Wood Chipper, Biochar Briket, Kompos Kering, Bahan RDF."],
        ["3", "KK", "Kertas / Karton", "Sampah selulosa lembaran kering yang dapat didaur ulang.", "Dus karton/gelombang, kertas HVS, majalah, koran, kemasan duplex.", "Pencacahan Daur Ulang Kertas, Bank Sampah, Upcycling."],
        ["4", "PL", "Plastik", "Sampah anorganik polimer sintetis yang lambat terurai.", "Botol PET, gelas PP, kantong kresek LDPE, wadah HDPE, kemasan saset.", "Pencacahan Pellet, Bank Sampah, Pirolisis BBM, Ecobricks."],
        ["5", "LG", "Logam", "Sampah anorganik logam ferrous & non-ferrous bernilai tinggi.", "Kaleng minuman alumunium, seng, paku, besi tua, kawat, tembaga.", "Bank Sampah, Peleburan Industri Daur Ulang Metal."],
        ["6", "KT", "Kain / Tekstil", "Sampah serat kain, pakaian bekas, dan sisa bahan konveksi.", "Pakaian bekas, perca garmen, sprei, handuk, benang, selimut.", "Bank Sampah (Upcycling Perca), Kain Majun Industri, Bahan RDF."],
        ["7", "KL", "Karet / Kulit", "Sampah anorganik elastomer karet dan kulit olahan hewan.", "Ban bekas, sandal jepit, sepatu bekas, tas kulit sintetis, sarung tangan.", "Upcycling Pot Ban, Furnitur Ban, Refuse Derived Fuel (RDF)."],
        ["8", "KC", "Kaca", "Sampah anorganik silika cair yang tidak mudah terdegradasi.", "Botol sirup/kecap kaca, jar stoples, pecahan kaca jendela, cermin.", "Penyetoran Botol Guna Ulang (Returnable Glass), Crusher Kaca."]
    ]

    table_sipsn = doc.add_table(rows=len(data_sipsn) + 1, cols=6)
    table_sipsn.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_sipsn = [Inches(0.4), Inches(0.6), Inches(1.3), Inches(1.5), Inches(1.5), Inches(1.5)]
    
    # Header Format
    hdr_cells = table_sipsn.rows[0].cells
    for i, title in enumerate(headers_sipsn):
        hdr_cells[i].text = title
        hdr_cells[i].width = widths_sipsn[i]
        set_cell_background(hdr_cells[i], HEX_PRIMARY)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Format
    for row_idx, row_data in enumerate(data_sipsn):
        row_cells = table_sipsn.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_sipsn[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=80, right=80)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8.5)
                if col_idx in [1, 2]:
                    run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # BAGIAN 2: KATEGORI SPESIAL SIMPAH
    add_heading_1(doc, "Bagian 2: Kode Kategori Tambahan di Aplikasi SIMPAH")
    add_body_p(doc, "Aplikasi SIMPAH melengkapi pencatatan lapangan dengan 2 kode kategori operasional tambahan:")

    headers_special = ["Kode System", "Nama Kategori", "Fungsi & Operasional Aplikasi SIMPAH"]
    data_special = [
        ["MIX", "Sampah Campuran", "Digunakan pada tahap Sampah Masuk sebelum dilakukan pemilahan (sorting) di TPS3R atau TPS desa."],
        ["LN", "Lainnya / Residu", "Digunakan untuk residu sampah yang tidak dapat didaur ulang/diolah lagi (popok bekas, B3 domestik) yang langsung dikirim ke TPA."]
    ]

    tbl_sp = doc.add_table(rows=len(data_special) + 1, cols=3)
    tbl_sp.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_sp = [Inches(1.2), Inches(1.8), Inches(3.8)]

    hdr_sp = tbl_sp.rows[0].cells
    for i, title in enumerate(headers_special):
        hdr_sp[i].text = title
        hdr_sp[i].width = widths_sp[i]
        set_cell_background(hdr_sp[i], HEX_HEADER_SPECIAL)
        set_cell_margins(hdr_sp[i], top=90, bottom=90, left=100, right=100)
        p = hdr_sp[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data_special):
        row_cells = tbl_sp.rows[row_idx + 1].cells
        bg_color = HEX_LIGHT_ROW if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            row_cells[col_idx].width = widths_sp[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=100, right=100)
            set_cell_border(row_cells[col_idx], 
                            bottom={'val': 'single', 'color': 'E5E7EB', 'sz': '4'},
                            top={'val': 'single', 'color': 'E5E7EB', 'sz': '4'})
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                if col_idx < 2:
                    run.font.bold = True

    # SAVE DOCUMENT
    target_out = DOC_PATH
    try:
        doc.save(target_out)
        print(f"Document successfully saved at: {target_out}")
    except PermissionError:
        target_out = DOC_FALLBACK_PATH
        doc.save(target_out)
        print(f"Primary locked. Document saved at fallback: {target_out}")

if __name__ == "__main__":
    build_document()
