import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    os.system(f'"{sys.executable}" -m pip install python-docx')
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = None # Reset color to black-ish

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('PROPOSAL LOMBA KRENOVA PROVINSI JAWA TENGAH 2026', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.bold = True

    add_paragraph(doc, '\n')
    
    # Identitas Inovasi
    add_heading(doc, 'A. IDENTITAS INOVASI', 1)
    add_paragraph(doc, '1. Judul Inovasi: SIMPAH (Sistem Informasi Manajemen Pengelolaan Sampah Terintegrasi AI)')
    add_paragraph(doc, '2. Kategori Lomba: [Isi dengan Kategori, misal: Tata Kelola Pemerintahan / Lingkungan Hidup]')
    add_paragraph(doc, '3. Nama Inventor / Tim: [Isi Nama Anda / Tim]')
    add_paragraph(doc, '4. Instansi / Asal Kabupaten: [Isi Asal Kabupaten/Kota]')
    add_paragraph(doc, '5. Kontak (HP/Email): [Isi Nomor HP dan Email]')
    
    add_paragraph(doc, '\n')
    
    # Latar Belakang
    add_heading(doc, 'B. LATAR BELAKANG & IDENTIFIKASI MASALAH', 1)
    p = add_paragraph(doc, 'Pengelolaan persampahan di tingkat daerah seringkali menghadapi kendala dalam hal akurasi data dan pengambilan keputusan yang tepat waktu. Pencatatan harian di tingkat Tempat Pembuangan Sementara (TPS) dan Tempat Pengolahan Sampah Terpadu (TPST) yang sebagian besar masih dilakukan secara manual menggunakan kertas, rentan terhadap kehilangan data dan manipulasi. Selain itu, banyak lokasi TPS berada di wilayah minim koneksi internet (blank spot), yang menyebabkan keterlambatan pelaporan data harian ke pusat tingkat Kabupaten/Kota.')
    p = add_paragraph(doc, 'Dampak dari tidak validnya data ini membuat Dinas Lingkungan Hidup kesulitan memprediksi secara akurat kebutuhan penjemputan armada truk sampah harian. Hal ini memicu masalah operasional, seperti penumpukan sampah liar akibat keterlambatan armada atau inefisiensi BBM karena armada mendatangi TPS yang volumenya masih sedikit. Proses rekapitulasi data untuk standar SIPSN (Sistem Informasi Pengelolaan Sampah Nasional) KLHK juga memakan waktu lama dan tidak efisien.')
    
    # Solusi Inovasi
    add_heading(doc, 'C. PROFIL INOVASI (SOLUSI)', 1)
    add_paragraph(doc, 'SIMPAH hadir sebagai platform digital terintegrasi untuk mengatasi masalah tersebut, membawa teknologi mutakhir hingga ke akar rumput. SIMPAH bukan sekadar aplikasi pencatatan, melainkan sebuah Sistem Cerdas (Smart System) dengan fitur unggulan:')
    
    add_paragraph(doc, '1. Progressive Web App (PWA) Offline-First: ', bold=True)
    doc.paragraphs[-1].add_run('Petugas lapangan dapat mengakses sistem langsung dari peramban (browser) handphone tanpa perlu menginstal aplikasi berat dari PlayStore. Sistem didesain untuk bisa menginput data meskipun sedang offline (tidak ada sinyal), dan akan secara otomatis melakukan sinkronisasi ketika perangkat kembali online.')
    
    add_paragraph(doc, '2. Predictive Analytics dengan Machine Learning (Prophet): ', bold=True)
    doc.paragraphs[-1].add_run('Mengintegrasikan algoritma Machine Learning Prophet (dikembangkan oleh Meta) untuk menganalisis data historis sampah yang fluktuatif. Sistem secara otomatis memproyeksikan (forecasting) volume timbulan sampah hingga 7 hari ke depan dengan mempertimbangkan tren mingguan dan harian.')
    
    add_paragraph(doc, '3. Asisten AI Edukasi (SIMPAH Buddy): ', bold=True)
    doc.paragraphs[-1].add_run('Dilengkapi dengan asisten cerdas berbasikan Large Language Model (Qwen) yang mampu merangkum laporan bagi para pembuat kebijakan (Eksekutif) secara interaktif, serta dapat mengedukasi warga tentang jenis dan pemilahan sampah dengan bahasa yang mudah dipahami.')
    
    # Penerapan Inovasi
    add_heading(doc, 'D. PENERAPAN & IMPLEMENTASI INOVASI', 1)
    add_paragraph(doc, 'Target pengguna inovasi ini meliputi 3 lapisan utama:')
    add_paragraph(doc, '- Petugas TPS/Kader Lingkungan: Sebagai ujung tombak penginputan data harian melalui smartphone.')
    add_paragraph(doc, '- Koordinator TPST/Kecamatan: Melakukan validasi data sebelum dikirimkan ke sistem pusat.')
    add_paragraph(doc, '- Kepala Dinas/Eksekutif: Mengakses Executive Dashboard untuk melihat KPI (Key Performance Indicator), grafik prediksi sampah, dan jumlah MoU armada yang aktif.')
    add_paragraph(doc, 'Saat ini inovasi SIMPAH [Isi dengan status saat ini, misal: telah dalam tahap penyelesaian prototipe dan siap dilakukan Uji Coba (UAT) di Kabupaten X].')
    
    # Manfaat Inovasi
    add_heading(doc, 'E. MANFAAT INOVASI & DAMPAK', 1)
    add_paragraph(doc, '1. Dampak Ekonomi: ', bold=True)
    doc.paragraphs[-1].add_run('Efisiensi anggaran operasional dan BBM armada pengangkut sampah hingga [Isi persentase estimasi, misal: 20%] karena penjemputan armada akan didasarkan pada data analitik prediksi (Data-driven fleet routing), bukan sekadar rute statis.')
    
    add_paragraph(doc, '2. Dampak Birokrasi: ', bold=True)
    doc.paragraphs[-1].add_run('Kesesuaian dengan format SIPSN memotong waktu berhari-hari kerja admin Dinas Lingkungan Hidup dalam merekap data komposisi sampah tahunan.')
    
    add_paragraph(doc, '3. Dampak Lingkungan: ', bold=True)
    doc.paragraphs[-1].add_run('Mencegah terjadinya penumpukan sampah yang membusuk di TPS karena prediktabilitas yang tinggi, serta meningkatkan angka pengurangan sampah lewat monitoring terpilah.')
    
    # Komersialisasi & Keberlanjutan
    add_heading(doc, 'F. RENCANA KOMERSIALISASI & KEBERLANJUTAN', 1)
    add_paragraph(doc, 'Inovasi ini dirancang agar dapat hidup secara mandiri (Sustain) melalui model bisnis:')
    add_paragraph(doc, '1. B2G (Business to Government): Penawaran lisensi perangkat lunak dan biaya pemeliharaan (maintenance) tahunan ke Dinas/Pemerintah Daerah setempat.')
    add_paragraph(doc, '2. B2B (Business to Business) / SaaS: Penawaran platform manajemen bagi pengelola kawasan perumahan mandiri, kawasan industri, dan hotel besar.')
    add_paragraph(doc, '3. Ke depannya, inovasi ini direncanakan akan diintegrasikan dengan Payment Gateway (seperti Midtrans) untuk sistem pemungutan retribusi persampahan warga secara digital, mencegah pungutan liar.')
    
    # Penutup
    add_paragraph(doc, '\n')
    add_paragraph(doc, 'Demikian proposal inovasi ini disusun. Data dan informasi dalam dokumen ini diklaim sebagai ide dan kreativitas orisinal tim pengembang.')
    
    # Signatures
    add_paragraph(doc, '\n')
    add_paragraph(doc, '[Kota, Tanggal Bulan Tahun]')
    add_paragraph(doc, 'Inventor / Perwakilan Tim,')
    add_paragraph(doc, '\n\n\n')
    add_paragraph(doc, '(.......................................................)')

    # Save
    doc.save(r'u:\Project\simpah-rilis v1\docs\Proposal_Krenova_SIMPAH.docx')
    print("Proposal saved successfully!")

if __name__ == '__main__':
    main()
